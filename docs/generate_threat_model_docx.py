#!/usr/bin/env python3
"""
Generate the Thunder OAuth 2.0 Threat Model DOCX by cloning the WSO2 template
and replacing/extending its content with the actual threat model data.

v2: Fixes formatting consistency, adds missing features from codebase audit,
    embeds Mermaid-rendered PNG diagrams.

Usage:
    python3 docs/generate_threat_model_docx.py

Output:
    docs/[Threat Model] - Thunder OAuth2.docx
"""

import copy
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "[Threat Model] - Template.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "[Threat Model] - Thunder OAuth2.docx")
DIAGRAMS_DIR = os.path.join(SCRIPT_DIR, "threat-model-diagrams")

# ─── Standard formatting constants (match template: Arial, docDefaults) ────────
BODY_SIZE = 11
SMALL_SIZE = 10
TABLE_HEADER_SIZE = 10
FONT_NAME = "Arial"

# Spacing constants (in EMU — 1 pt = 12700 EMU)
SPACE_AFTER_BODY = Pt(6)       # 6pt after body paragraphs
SPACE_BEFORE_LABEL = Pt(10)    # 10pt before bold sub-labels
SPACE_AFTER_LABEL = Pt(2)      # 2pt after bold sub-labels (tight to content)
SPACE_BEFORE_TABLE = Pt(4)     # 4pt before tables
SPACE_AFTER_TABLE = Pt(10)     # 10pt after tables
SPACE_AFTER_IMAGE = Pt(8)      # 8pt after images
LINE_SPACING = 1.15            # 1.15x line spacing for readability
CELL_MARGIN = Pt(3)            # 3pt cell top/bottom padding


# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_text(cell, text, bold=False, size=SMALL_SIZE, color=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fix_table_grid_widths(table):
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn('w:gridCol')):
            w = gridCol.get(qn('w:w'))
            if w is not None:
                try:
                    int(w)
                except ValueError:
                    gridCol.set(qn('w:w'), str(int(float(w))))


def add_table_borders(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_cell_margins(table, top=45, bottom=45, left=80, right=80):
    """Set default cell margins for a table (in twips: 1pt = 20 twips)."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    margins = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(margins)


def add_spacing_before_table(table, space_pt=4):
    """Add a spacer paragraph before a table element in the document body."""
    spacer = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:before="0" w:after="0" w:line="60" w:lineRule="exact"/>'
        f'  <w:rPr><w:sz w:val="2"/></w:rPr></w:pPr>'
        f'</w:p>'
    )
    table._tbl.addprevious(spacer)


def add_spacing_after_table(table, space_pt=10):
    """Add a spacer paragraph after a table element in the document body."""
    spacer = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:before="{space_pt * 20}" w:after="0" w:line="240" w:lineRule="auto"/>'
        f'  <w:rPr><w:sz w:val="2"/></w:rPr></w:pPr>'
        f'</w:p>'
    )
    table._tbl.addnext(spacer)


def add_row(table, cells_data, bold=False, size=SMALL_SIZE, header=False):
    fix_table_grid_widths(table)
    row = table.add_row()
    for i, text in enumerate(cells_data):
        if i < len(row.cells):
            set_cell_text(row.cells[i], str(text), bold=bold, size=size)
            if header:
                set_cell_shading(row.cells[i], "D9E2F3")
    return row


def clear_table_data_rows(table, keep_header=True):
    start = 1 if keep_header else 0
    while len(table.rows) > start:
        tr = table.rows[len(table.rows) - 1]._tr
        table._tbl.remove(tr)


def find_paragraph_index(doc, text_start, start_from=0):
    for i in range(start_from, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith(text_start):
            return i
    return -1


def clear_paragraph(p):
    for r in p.runs:
        r.text = ""
    for elem in p._p.findall(qn('w:r')):
        for t in elem.findall(qn('w:t')):
            t.text = ""


def set_paragraph_text(p, text, bold=False, size=None):
    clear_paragraph(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size or BODY_SIZE)
    run.font.name = FONT_NAME
    p.paragraph_format.space_after = SPACE_AFTER_BODY
    p.paragraph_format.line_spacing = LINE_SPACING


def apply_label_spacing(p):
    """Apply label-style spacing (extra before, tight after) to an existing paragraph."""
    p.paragraph_format.space_before = SPACE_BEFORE_LABEL
    p.paragraph_format.space_after = SPACE_AFTER_LABEL
    p.paragraph_format.line_spacing = LINE_SPACING


def insert_paragraph_after_p(doc, paragraph, text, bold=False, size=BODY_SIZE):
    new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = "normal"
    new_para.paragraph_format.space_after = SPACE_AFTER_BODY
    new_para.paragraph_format.line_spacing = LINE_SPACING
    run = new_para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    return new_para


def insert_image_after(doc, paragraph, image_path, width_inches=6.0):
    if not os.path.exists(image_path):
        print(f"  Warning: Image not found: {image_path}")
        return None
    new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    img_para = Paragraph(new_p, paragraph._parent)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = SPACE_AFTER_IMAGE
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return img_para


def insert_image_before(insert_point, doc, image_path, width_inches=6.0):
    if not os.path.exists(image_path):
        print(f"  Warning: Image not found: {image_path}")
        return
    new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
    insert_point.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    img_para = Paragraph(new_p, doc.paragraphs[0]._parent)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = SPACE_AFTER_IMAGE
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))


def remove_purple_paragraphs(doc):
    """Remove template instruction paragraphs (purple-colored text, Note: guidance)."""
    to_remove = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        for r in p.runs:
            if r.font.color and r.font.color.rgb:
                rgb = str(r.font.color.rgb)
                if rgb in ("7030A0", "800080", "9900FF", "663399"):
                    to_remove.append(p._p)
                    break
        if text.startswith("Note:") and any(kw in text for kw in [
            "should showcase", "helpful for the reviewers", "doing interaction-level",
            "leveraging any existing", "following is an example",
            "complex use cases", "When you have complex"
        ]):
            to_remove.append(p._p)
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)


class InteractionBuilder:
    """Builds interaction sections before a given insert point."""

    def __init__(self, doc, insert_point):
        self.doc = doc
        self.insert_point = insert_point

    def _add_heading(self, text, level=3):
        new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
        self.insert_point.addprevious(new_p)
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_p, self.doc.paragraphs[0]._parent)
        para.style = f"Heading {level}"
        run = para.add_run(text)
        run.font.name = FONT_NAME
        # Headings get their spacing from style definition (H3: 16pt before, 4pt after)
        return para

    def _add_text(self, text, bold=False, size=BODY_SIZE, is_label=False):
        new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
        self.insert_point.addprevious(new_p)
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_p, self.doc.paragraphs[0]._parent)
        para.style = "normal"
        if is_label:
            # Sub-labels like "Description", "Assets Involved" etc.
            para.paragraph_format.space_before = SPACE_BEFORE_LABEL
            para.paragraph_format.space_after = SPACE_AFTER_LABEL
        else:
            para.paragraph_format.space_after = SPACE_AFTER_BODY
        para.paragraph_format.line_spacing = LINE_SPACING
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = FONT_NAME
        return para

    def _add_table(self, cols, header, data):
        table = self.doc.add_table(rows=1, cols=cols)
        add_table_borders(table)
        set_table_cell_margins(table)
        for i, h in enumerate(header):
            set_cell_text(table.rows[0].cells[i], h, bold=True, size=TABLE_HEADER_SIZE)
            set_cell_shading(table.rows[0].cells[i], "D9E2F3")
        for row_data in data:
            add_row(table, row_data, size=SMALL_SIZE)
        self.insert_point.addprevious(table._tbl)
        return table

    def _add_image(self, image_path, width_inches=5.8):
        insert_image_before(self.insert_point, self.doc, image_path, width_inches)

    def _add_spacer(self):
        self._add_text("")

    def build_interaction(self, iid, title, description,
                          assets_data, data_flow_text, access_control_text,
                          security_data, threats_data, diagram_path=None):
        self._add_heading(f"[{iid}]: {title}")
        self._add_text("Description", bold=True, is_label=True)
        self._add_text(description)

        if diagram_path:
            self._add_text("Data Flow Diagram", bold=True, is_label=True)
            self._add_image(diagram_path)

        self._add_text("Assets Involved", bold=True, is_label=True)
        self._add_table(3, ["Initiator", "Intermediate", "Target"], assets_data)
        self._add_text("Data Flow", bold=True, is_label=True)
        self._add_text(data_flow_text)
        self._add_text("Access Control", bold=True, is_label=True)
        self._add_text(access_control_text)
        self._add_text("Security Considerations", bold=True, is_label=True)
        self._add_table(3, ["Area", "Response", "Comments"], security_data)
        self._add_text("Threat Assessment", bold=True, is_label=True)
        self._add_table(5,
            ["ID", "Category", "Threat", "Materializable", "Mitigations / Comment"],
            threats_data)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("Loading template...")
    doc = Document(TEMPLATE_PATH)

    # ── 0. Remove purple instruction text ──
    print("  Removing template instructions...")
    remove_purple_paragraphs(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 1. Cover Page
    # ═════════════════════════════════════════════════════════════════════
    print("  Filling cover page...")
    for p in doc.paragraphs:
        if "<Threat Model Name>" in p.text:
            clear_paragraph(p)
            run = p.add_run("Thunder OAuth 2.0 / OIDC Component")
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = FONT_NAME
            break

    # ═════════════════════════════════════════════════════════════════════
    # 2. Revision History (Table 0)
    # ═════════════════════════════════════════════════════════════════════
    t0 = doc.tables[0]
    set_table_cell_margins(t0)
    set_cell_text(t0.rows[1].cells[0], "v1.0")
    set_cell_text(t0.rows[1].cells[1], "2025-02-24")
    set_cell_text(t0.rows[1].cells[2], "Thunder Team")
    set_cell_text(t0.rows[1].cells[3], "Initial version — Complete OAuth 2.0/OIDC threat model")

    # ═════════════════════════════════════════════════════════════════════
    # 3. Introduction
    # ═════════════════════════════════════════════════════════════════════
    print("  Writing introduction...")
    idx = find_paragraph_index(doc, "[Provide an overview")
    if idx >= 0:
        set_paragraph_text(doc.paragraphs[idx],
            "Thunder is a lightweight user and identity management product providing "
            "authentication and authorization capabilities. This threat model covers "
            "the OAuth 2.0 / OpenID Connect (OIDC) authorization server component.\n\n"
            "Scope:\n"
            "\u2022 Authorization Code Grant with PKCE (RFC 7636)\n"
            "\u2022 Client Credentials Grant\n"
            "\u2022 Refresh Token Grant (with configurable rotation)\n"
            "\u2022 Token Exchange Grant (RFC 8693) — delegation/impersonation with act claim\n"
            "\u2022 Token Introspection (RFC 7662)\n"
            "\u2022 OIDC UserInfo Endpoint (live user data)\n"
            "\u2022 JWKS Endpoint (RSA, ECDSA P-256/P-384/P-521, EdDSA Ed25519)\n"
            "\u2022 OAuth 2.0 / OIDC Discovery (RFC 8414)\n"
            "\u2022 Dynamic Client Registration (RFC 7591)\n"
            "\u2022 Client Authentication (client_secret_basic, client_secret_post, none)\n"
            "\u2022 OIDC claims request parameter\n"
            "\u2022 Resource Indicator (RFC 8707)\n"
            "\u2022 Per-application token, claim, and scope configuration\n"
            "\u2022 Observability events for token lifecycle\n\n"
            "Standards: RFC 6749, 7636, 7662, 8414, 8693, 8707, OIDC Core 1.0, "
            "RFC 9700 (OAuth 2.0 Security BCP, Jan 2025).\n\n"
            "Not yet implemented: Token Revocation (/oauth2/revoke), "
            "Logout (/oauth2/logout), Consent screen (TODO in codebase).")

    idx2 = find_paragraph_index(doc, "[Attach any associated")
    if idx2 >= 0:
        set_paragraph_text(doc.paragraphs[idx2],
            "Associated docs: api/authentication.yaml, api/flow-execution.yaml, "
            "api/application.yaml, api/design.yaml")

    # ═════════════════════════════════════════════════════════════════════
    # 4. Architecture Diagram
    # ═════════════════════════════════════════════════════════════════════
    print("  Adding architecture diagram...")
    idx = find_paragraph_index(doc, "[Architecture diagrams")
    if idx < 0:
        idx = find_paragraph_index(doc, "Architecture Diagram (text representation")
    if idx >= 0:
        set_paragraph_text(doc.paragraphs[idx],
            "The following diagram shows the high-level architecture of the Thunder "
            "OAuth 2.0 / OIDC module, including endpoints, internal components, "
            "dependencies, and data stores.")
        insert_image_after(doc, doc.paragraphs[idx],
                           os.path.join(DIAGRAMS_DIR, "architecture.png"), 6.2)

    # ═════════════════════════════════════════════════════════════════════
    # 5. Data Flow Diagrams
    # ═════════════════════════════════════════════════════════════════════
    print("  Adding data flow diagrams...")
    idx = find_paragraph_index(doc, "[Data Flow Diagrams")
    if idx < 0:
        idx = find_paragraph_index(doc, "Key Data Flows")
    if idx >= 0:
        set_paragraph_text(doc.paragraphs[idx],
            "The following sequence diagrams illustrate the key data flows.")
        # Insert in reverse order (each inserts AFTER idx paragraph)
        flows = [
            ("DF-4: Token Exchange (RFC 8693)", "flow-token-exchange.png"),
            ("DF-3: Refresh Token Grant", "flow-refresh-token.png"),
            ("DF-2: Client Credentials Grant", "flow-client-credentials.png"),
            ("DF-1: Authorization Code Grant with PKCE", "flow-authcode-pkce.png"),
        ]
        for title, filename in flows:
            img_path = os.path.join(DIAGRAMS_DIR, filename)
            title_p = insert_paragraph_after_p(doc, doc.paragraphs[idx], title,
                                               bold=True, size=BODY_SIZE)
            insert_image_after(doc, title_p, img_path, 6.0)

    # ═════════════════════════════════════════════════════════════════════
    # 6. Actors Table (T2)
    # ═════════════════════════════════════════════════════════════════════
    print("  Filling actors, resources, deps, trust boundaries...")
    t2 = doc.tables[2]
    set_table_cell_margins(t2)
    clear_table_data_rows(t2)
    actors = [
        ("End User (Resource Owner)",
         "Authenticates via browser to authorize client access to their resources.",
         "Initiate auth flow, authenticate, authorize/deny, access UserInfo"),
        ("OAuth Client (Confidential)",
         "Server-side app with client secret. Can securely store credentials.",
         "Auth code, client credentials, refresh, token exchange grants"),
        ("OAuth Client (Public)",
         "SPA/native app that cannot securely store secrets.",
         "Auth code + PKCE, refresh token grants"),
        ("Resource Server",
         "API server validating access tokens to protect resources.",
         "Introspect tokens, validate JWT via JWKS"),
        ("Admin / Developer",
         "Manages OAuth apps and configuration.",
         "Register apps (DCR/mgmt API), configure grants, redirect URIs, token policies"),
        ("Web Attacker (A1, RFC 9700 Section 3)",
         "Controls web servers. Creates redirect URIs, phishes users, manipulates browser flows.",
         "Code interception, CSRF, phishing, open redirect, mix-up attacks"),
        ("Network Attacker (A2, RFC 9700 Section 3)",
         "Full network control: eavesdrop, intercept, modify, replay. Assumes TLS not broken.",
         "MITM (if TLS misconfigured), token interception, TLS downgrade"),
    ]
    for a in actors:
        add_row(t2, a)

    # ═════════════════════════════════════════════════════════════════════
    # 7. Entitlement Matrix (T3)
    # ═════════════════════════════════════════════════════════════════════
    t3 = doc.tables[3]
    set_table_cell_margins(t3)
    clear_table_data_rows(t3)
    h3 = t3.rows[0]
    ent_h = ["Actor", "Auth Flow", "Token Exchange", "Introspect", "UserInfo"]
    for i, h in enumerate(ent_h[:len(h3.cells)]):
        set_cell_text(h3.cells[i], h, bold=True, size=TABLE_HEADER_SIZE)
    ent_data = [
        ("End User", "Yes (browser)", "No", "No", "No"),
        ("Confidential Client", "Yes + secret", "Yes", "Yes*", "Yes (Bearer)"),
        ("Public Client", "Yes + PKCE", "No", "Yes*", "Yes (Bearer)"),
        ("Resource Server", "No", "No", "Yes*", "No"),
        ("Admin", "No", "No", "No", "No"),
    ]
    for e in ent_data:
        add_row(t3, e[:len(h3.cells)])

    # ═════════════════════════════════════════════════════════════════════
    # 8. Resources (T4)
    # ═════════════════════════════════════════════════════════════════════
    t4 = doc.tables[4]
    set_table_cell_margins(t4)
    clear_table_data_rows(t4)
    resources = [
        ("Authorization Endpoint",
         "GET /oauth2/authorize \u2014 Validates client, redirect URI (exact match), PKCE, resource (RFC 8707), "
         "claims param; stores auth request context (10-min, single-use) in runtime DB."),
        ("Auth Callback",
         "POST /oauth2/auth/callback \u2014 Verifies signed JWT assertion from Flow Engine; generates UUIDv7 "
         "auth code; validates sub claim constraints; redirects with 302."),
        ("Token Endpoint",
         "POST /oauth2/token \u2014 Client Auth Middleware (Basic/Post/None, constant-time SHA-256). "
         "4 grant types. Cache-Control: no-store, Pragma: no-cache. Publishes observability events."),
        ("Introspection",
         "POST /oauth2/introspect \u2014 Validates tokens. GAP: No client auth. "
         "TODO: Token revocation/validity checks not implemented."),
        ("UserInfo",
         "GET|POST /oauth2/userinfo \u2014 Bearer token required (openid scope). Fetches LIVE user data "
         "(attributes, groups \u226420, ouId, ouName). Rejects client_credentials tokens. Cache-Control: no-store."),
        ("JWKS",
         "GET /oauth2/jwks \u2014 RSA, ECDSA (P-256/P-384/P-521), EdDSA (Ed25519). "
         "x5c chain, x5t, x5t#S256 thumbprints."),
        ("Discovery",
         "/.well-known/openid-configuration + /.well-known/oauth-authorization-server \u2014 RFC 8414 / OIDC Discovery."),
        ("DCR",
         "POST /oauth2/dcr/register \u2014 RFC 7591. jwks_uri/jwks (mutually exclusive). "
         "Auth: dcr.insecure=true \u2192 open; false \u2192 system permission. "
         "client_secret_expires_at=0 (never)."),
        ("Authorization Codes",
         "UUIDv7, single-use, configurable validity. Stored with PKCE challenge, redirect URI, scope, "
         "state, resource. Deactivated on first retrieval."),
        ("Auth Request Contexts",
         "Ephemeral request data. 10-min hardcoded expiry. Single-use (deleted on read). Contains all OAuth params."),
        ("Access Tokens",
         "JWT/RS256. Claims: sub, aud, iss, scope, client_id, grant_type, claims_req, claims_locales, "
         "user attributes, act (delegation). Per-app validity and user_attributes."),
        ("Refresh Tokens",
         "JWT with access_token_sub, access_token_aud, metadata for re-derivation. "
         "Rotation via renew_on_grant config."),
        ("ID Tokens",
         "OIDC JWT with auth_time, standard claims per scope, per-app filtered."),
        ("PKI Signing Keys",
         "RSA/ECDSA/EdDSA private keys on filesystem. Multiple keys; preferred_key_id selects active."),
        ("Client Secrets",
         "SHA-256 hashed. Constant-time comparison (crypto/subtle). Never expire (client_secret_expires_at=0)."),
        ("Observability Events",
         "TOKEN_ISSUANCE_STARTED/ISSUED/FAILED with client_id, grant_type, scope, error, duration_ms."),
        ("OIDC Claims/Scopes",
         "Standard: openid, profile, email, phone, address. Per-app scope_claims override. "
         "claims parameter with essential/value/values. Custom: userType, ouId, ouName, ouHandle, groups."),
    ]
    for r in resources:
        add_row(t4, r)

    # ═════════════════════════════════════════════════════════════════════
    # 9. Dependencies (T5)
    # ═════════════════════════════════════════════════════════════════════
    t5 = doc.tables[5]
    set_table_cell_margins(t5)
    clear_table_data_rows(t5)
    deps = [
        ("Application Service",
         "internal/application \u2014 OAuth app config, credentials (SHA-256 constant-time), redirect URIs, "
         "grants, auth methods, token config, user_attributes, scope_claims."),
        ("Flow Engine",
         "internal/flow/flowexec \u2014 User authentication (password, passwordless, MFA, social). "
         "Returns signed JWT assertion with identity and auth_time."),
        ("User Service",
         "internal/user \u2014 User attributes (JSON), groups (\u226420 via Group Service), "
         "ouId, ouName, ouHandle, userType."),
        ("PKI Service",
         "internal/system/crypto/pki \u2014 X.509 certs and private keys (RSA/ECDSA/EdDSA)."),
        ("JWT Service",
         "internal/system/jose/jwt \u2014 JWT generation and verification. Configurable issuer, audience, leeway."),
        ("Database",
         "internal/system/database \u2014 SQLite (dev) or PostgreSQL (prod). "
         "DEPLOYMENT_ID for multi-tenant isolation."),
        ("TLS",
         "TLS 1.3 minimum in deployment.yaml. Server cert + key on filesystem."),
        ("Gate UI",
         "frontend/apps/thunder-gate \u2014 Login/register/recovery UI."),
        ("Middleware",
         "internal/system/middleware \u2014 CORS on all endpoints. security.GetPermissions for DCR auth."),
    ]
    for d in deps:
        add_row(t5, d)

    # ═════════════════════════════════════════════════════════════════════
    # 10. Trust Boundaries (T6) + diagram
    # ═════════════════════════════════════════════════════════════════════
    tb_idx = find_paragraph_index(doc, "This section aims to identify")
    if tb_idx >= 0:
        set_paragraph_text(doc.paragraphs[tb_idx],
            "Trust boundaries of the OAuth/OIDC component:")
        insert_image_after(doc, doc.paragraphs[tb_idx],
                           os.path.join(DIAGRAMS_DIR, "trust-boundaries.png"), 6.0)

    t6 = doc.tables[6]
    set_table_cell_margins(t6)
    clear_table_data_rows(t6)
    boundaries = [
        ("TB-01", "Untrust\u2192Trust", "Browser \u2192 GET /oauth2/authorize (auth request over TLS)"),
        ("TB-02", "Untrust\u2192Trust", "Client \u2192 POST /oauth2/token (with client auth over TLS)"),
        ("TB-03", "Untrust\u2192Trust", "RS/Client \u2192 POST /oauth2/introspect (over TLS)"),
        ("TB-04", "Untrust\u2192Trust", "External \u2192 POST /oauth2/dcr/register (over TLS)"),
        ("TB-05", "Untrust\u2192Trust", "External \u2192 GET /oauth2/jwks (public keys over TLS)"),
        ("TB-06", "Untrust\u2192Trust", "Client \u2192 GET|POST /oauth2/userinfo (Bearer token over TLS)"),
        ("TB-07", "Trust\u2192Trust", "OAuth \u2192 Application Service (in-process, validates client config)"),
        ("TB-08", "Trust\u2192Trust", "OAuth \u2192 User Service (in-process, user attributes/groups)"),
        ("TB-09", "Trust\u2192Trust", "OAuth \u2192 Flow Engine (in-process, user authentication)"),
        ("TB-10", "Internal", "OAuth \u2194 Runtime DB (auth codes, contexts; parameterized queries)"),
        ("TB-11", "Internal", "Token Service \u2192 PKI keys (filesystem, restricted permissions)"),
        ("TB-12", "Trust\u2192Untrust", "Thunder \u2192 Gate UI: HTTP 200 JSON {redirect_uri} with auth code. Gate UI performs client-side redirect."),
        ("TB-13", "Untrust\u2192Trust", "Gate UI \u2192 POST /oauth2/auth/callback (JWT assertion over TLS)"),
        ("TB-14", "Internal", "Token Service \u2192 event system (observability events)"),
    ]
    for b in boundaries:
        add_row(t6, b)

    # ═════════════════════════════════════════════════════════════════════
    # 11. Inherited / Out-of-Scope
    # ═════════════════════════════════════════════════════════════════════
    idx = find_paragraph_index(doc, "[inherited risk 1]")
    if idx >= 0:
        set_paragraph_text(doc.paragraphs[idx],
            "\u2022 OS, container runtime, physical infrastructure threats.\n"
            "\u2022 Flow Engine authentication mechanisms (separate threat model).\n"
            "\u2022 Management API for application CRUD (separate).\n"
            "\u2022 Compromised/rogue admin with direct DB access.\n"
            "\u2022 L3/L4 DDoS.\n"
            "\u2022 EOL software components in Go runtime or OS.\n"
            "\u2022 Gate UI frontend (XSS, CSRF on login \u2014 separate).\n"
            "\u2022 Consent flow threats (not yet implemented \u2014 TODO).\n"
            "\u2022 Token revocation threats (/oauth2/revoke \u2014 not implemented).\n"
            "\u2022 Logout/session management (/oauth2/logout \u2014 not implemented).\n"
            "\u2022 PAR, Device Authorization Grant (not implemented).")

    # ═════════════════════════════════════════════════════════════════════
    # 12. Template interactions I-01, I-02
    # ═════════════════════════════════════════════════════════════════════
    print("  Building interaction sections...")

    # ── I-01 ──
    idx = find_paragraph_index(doc, "[XX]: [Interaction Name]")
    if idx >= 0:
        set_paragraph_text(doc.paragraphs[idx],
            "[I-01]: Authorization Request (End User \u2192 Authorization Endpoint)",
            bold=True, size=14)

    idx_desc = find_paragraph_index(doc, "<Description of the interaction>")
    if idx_desc >= 0:
        set_paragraph_text(doc.paragraphs[idx_desc],
            "Browser GET to /oauth2/authorize with: response_type=code, client_id, redirect_uri, "
            "scope, state, code_challenge/method, resource (RFC 8707), claims, claims_locales (OIDC).\n\n"
            "Processing flow:\n"
            "1. Validates client_id existence; returns Gate error page if missing or unknown.\n"
            "2. Retrieves OAuth app via ApplicationService.GetOAuthApplication(client_id).\n"
            "3. Validates redirect_uri (exact match against registered URIs); error page if invalid.\n"
            "4. Validates response_type=code and authorization_code grant allowed for app.\n"
            "5. PKCE validation: code_challenge required for public clients (S256 or plain).\n"
            "6. Resource parameter validated (absolute URI, no fragment) per RFC 8707.\n"
            "7. Parses OIDC claims parameter (JSON with value/values mutual exclusivity).\n"
            "8. Separates OIDC vs non-OIDC scopes; computes required user attributes.\n"
            "9. Initiates auth flow via FlowExecService.InitiateFlow().\n"
            "10. Stores auth request context in AUTHORIZATION_REQUEST table (10-min TTL, single-use).\n"
            "11. Redirects (302) to Gate login page with authId, applicationId, flowId.\n\n"
            "Error branching: Pre-redirect_uri errors (bad client_id, bad redirect_uri) go to Gate error page. "
            "Post-redirect_uri errors redirect to client redirect_uri with error + error_description + state. "
            "HTTP redirect_uri gets showInsecureWarning=true flag.")

    t7 = doc.tables[7]
    set_table_cell_margins(t7)
    set_cell_text(t7.rows[1].cells[0], "End User (Browser)")
    set_cell_text(t7.rows[1].cells[1], "OAuth Client (redirect)")
    set_cell_text(t7.rows[1].cells[2], "Auth Endpoint \u2192 App Service \u2192 Runtime DB")

    idx_df = find_paragraph_index(doc, "<add a hyperlink to relevant diagram")
    if idx_df >= 0:
        set_paragraph_text(doc.paragraphs[idx_df], "See DF-1: Authorization Code Grant with PKCE, steps 1-6.")

    idx_ac = find_paragraph_index(doc, "<How access control is done")
    if idx_ac >= 0:
        set_paragraph_text(doc.paragraphs[idx_ac],
            "No auth (public). Validates: client_id, redirect_uri exact match, response_type=code, "
            "PKCE required for public clients, resource as absolute URI w/o fragment.")

    t8 = doc.tables[8]
    set_table_cell_margins(t8)
    for i, (area, resp) in enumerate([
        ("Data Confidentiality", "High [C-High]"),
        ("Communication Medium", "Network [M-NT]"),
        ("Transport Security", "TLS 1.3"),
        ("Authentication", "None (public)"),
        ("Accessibility", "Public"),
    ]):
        if i+1 < len(t8.rows):
            set_cell_text(t8.rows[i+1].cells[0], area)
            set_cell_text(t8.rows[i+1].cells[1], resp)
    set_cell_text(t8.rows[1].cells[2],
        "Params: client_id, redirect_uri, scope, PKCE, state, resource, claims. "
        "Leakage enables code injection/CSRF.")

    t9 = doc.tables[9]
    set_table_cell_margins(t9)
    clear_table_data_rows(t9)
    for t_data in [
        ("1","Spoofing","Redirect URI Manipulation (RFC 9700 Section 4.1): Redirects code to attacker.",
         "No","Exact match. No wildcards. Re-verified at token exchange."),
        ("2","Tampering","Code Injection (RFC 9700 Section 4.5): Stolen code injected into client callback.",
         "No","PKCE enforced for public clients. Confidential clients need secret."),
        ("3","Info Disc.","Code Leakage via Referer (RFC 9700 Section 4.2).",
         "No","Single-use codes. PKCE binding."),
        ("4","Info Disc.","Code in Browser History (RFC 9700 Section 4.3).",
         "No","Single-use, short-lived. PKCE verifier not in URL."),
        ("5","Spoofing","CSRF (RFC 9700 Section 4.7): Attacker initiates flow for victim.",
         "No","State param + PKCE = CSRF protection per RFC 9700 Section 2.1."),
        ("6","Spoofing","PKCE Downgrade (RFC 9700 Section 4.8): Strip code_challenge.",
         "No","Required for public clients. Stored challenge requires verifier."),
        ("7","Info Disc.","307 Redirect (RFC 9700 Section 4.12): Forwards creds.",
         "No","Uses 302 only."),
        ("8","Tampering","Mix-Up (RFC 9700 Section 4.4): Client sends code to wrong AS.",
         "No","Server metadata via /.well-known. iss claim in tokens."),
        ("9","Elevation","Open Redirector (RFC 9700 Section 4.11.2).",
         "No","Invalid URI \u2192 error page, not redirect."),
        ("10","Info Disc.","Auth context in DB.",
         "No","10-min expiry, single-use, UUID authId."),
        ("11","Spoofing","Redirect URI auto-select when omitted.",
         "No","Only when 1 URI registered (per OAuth 2.1)."),
        ("12","Info Disc.","HTTP redirect_uri allows code interception.",
         "Partial","showInsecureWarning flag set but http still allowed. Should reject http:// entirely per RFC 9700 Section 2.6. RISK REGISTRY."),
        ("13","Tampering","resource param mismatch between auth and token.",
         "No","Stored + re-validated at exchange. Absolute URI, no fragment."),
        ("14","Info Disc.","claims param requests excessive PII.",
         "No","Filtered by per-app user_attributes. sub constraint validated. essential enforcement TODO."),
        ("15","Spoofing","Clickjacking on authorization endpoint (RFC 9700 Section 4.16).",
         "Yes","RISK: No X-Frame-Options or CSP frame-ancestors header. RISK REGISTRY."),
        ("16","Info Disc.","Missing Referrer-Policy header (RFC 9700 Section 4.2.4).",
         "Yes","RISK: No Referrer-Policy header on authorization responses. RISK REGISTRY."),
        ("17","Tampering","state parameter: no length or charset validation.",
         "Partial","State is stored in DB and included in redirect URIs. No length limit could cause large DB entries. URL encoding provides some protection but no input validation."),
        ("18","Spoofing","PKCE plain method accepted at authorization endpoint (RFC 9700 Section 2.1.1).",
         "Yes","RISK: plain method is accepted and advertised in metadata. Allows code_challenge == code_verifier with no hashing. S256 should be mandatory. See I-14 T6. RISK REGISTRY."),
        ("19","DoS","Authorization request flooding: no rate limit on /oauth2/authorize.",
         "Yes","RISK: Each authorize request creates a new auth request in DB and initiates a flow. No rate limiting. "
         "Combined with no DB cleanup (I-10 T8), enables resource exhaustion. RISK REGISTRY."),
    ]:
        add_row(t9, t_data)

    # ── I-02 ──
    idx2 = find_paragraph_index(doc, "[XX]: [Interaction Name]")
    if idx2 >= 0:
        set_paragraph_text(doc.paragraphs[idx2],
            "[I-02]: Authorization Callback (Flow Engine \u2192 Auth Callback)",
            bold=True, size=14)

    idx_d2 = find_paragraph_index(doc, "<Description of the interaction>")
    if idx_d2 >= 0:
        set_paragraph_text(doc.paragraphs[idx_d2],
            "Gate UI sends POST /oauth2/auth/callback with JSON body {authId, assertion}.\n\n"
            "Processing flow:\n"
            "1. Loads auth request context from AUTHORIZATION_REQUEST table (SELECT with expiry check, "
            "then DELETE — single-use, 10-min TTL enforced at query level).\n"
            "2. Verifies JWT assertion signature via jwtService.VerifyJWT().\n"
            "3. Decodes assertion claims: extracts sub, authorized_permissions, iat (becomes auth_time), "
            "and user attributes (excluding standard JWT claims).\n"
            "4. Validates sub claim constraint if openid scope present (OIDC claims parameter sub matching).\n"
            "5. Overwrites non-OIDC scopes with authorized_permissions from assertion.\n"
            "6. Creates authorization code: generates UUIDv7 for code ID and code value, sets validity "
            "from OAuth.AuthorizationCode.ValidityPeriod config.\n"
            "7. Persists auth code in AUTHORIZATION_CODE table with: PKCE challenge/method, redirect_uri, "
            "scope, state, resource, claims, claims_locales, user attributes.\n"
            "8. Returns HTTP 200 JSON {redirect_uri: 'client-redirect?code=xxx&state=yyy'}.\n\n"
            "IMPORTANT: The callback does NOT issue an HTTP 302 redirect. It returns JSON to Gate UI, "
            "which performs the client-side redirect via JavaScript.")

    t10 = doc.tables[10]
    set_table_cell_margins(t10)
    set_cell_text(t10.rows[1].cells[0], "Gate UI (Browser)")
    set_cell_text(t10.rows[1].cells[1], "Flow Engine (JWT)")
    set_cell_text(t10.rows[1].cells[2], "Auth Callback \u2192 Runtime DB")

    idx_df2 = find_paragraph_index(doc, "<add a hyperlink to relevant diagram")
    if idx_df2 >= 0:
        set_paragraph_text(doc.paragraphs[idx_df2], "See DF-1, steps 7-10.")
    idx_ac2 = find_paragraph_index(doc, "<How access control is done")
    if idx_ac2 >= 0:
        set_paragraph_text(doc.paragraphs[idx_ac2],
            "JWT assertion signed by server RSA key, verified via jwtService. "
            "Context single-use prevents replay.")

    t11 = doc.tables[11]
    set_table_cell_margins(t11)
    for i, (area, resp) in enumerate([
        ("Data Confidentiality", "High [C-High]"),
        ("Communication Medium", "Network [M-NT]"),
        ("Transport Security", "TLS 1.3"),
        ("Authentication", "JWT Assertion"),
        ("Accessibility", "Public"),
    ]):
        if i+1 < len(t11.rows):
            set_cell_text(t11.rows[i+1].cells[0], area)
            set_cell_text(t11.rows[i+1].cells[1], resp)
    set_cell_text(t11.rows[1].cells[2],
        "JWT assertion with user identity. Generates UUIDv7 code. TLS 1.3. Internal PKI signed.")

    t12 = doc.tables[12]
    set_table_cell_margins(t12)
    clear_table_data_rows(t12)
    for t_data in [
        ("1","Spoofing","Forged JWT Assertion.",
         "No","Signed with server PKI key. Verified cryptographically via jwtService."),
        ("2","Replay","JWT Assertion Replay.",
         "No","Context single-use (deleted on read). 10-min expiry enforced at DB query level."),
        ("3","Tampering","Sub Claim Bypass.",
         "No","OIDC claims sub constraint validated against assertion sub."),
        ("4","Spoofing","AuthId Guessing.",
         "No","UUID, single-use, 10-min expiry. Deleted immediately after retrieval."),
        ("5","Info Disc.","Auth code predictability (UUIDv7).",
         "Partial","UUIDv7 has 48-bit timestamp + 74 random bits. Less entropy than a pure random token. RFC 6749 Section 10.10 recommends high entropy. 74 bits is non-trivial but weaker than recommended 128+ bits."),
        ("6","Tampering","Authorized permissions escalation.",
         "No","Scopes overwritten with assertion authorized_permissions, not merged."),
        ("7","DoS","Auth context consumed before assertion verification.",
         "Yes","RISK: loadAuthRequestContext deletes the auth request from the store BEFORE verifyAssertion is called. "
         "An attacker who obtains the authId can call the callback with an invalid assertion, consuming the auth context "
         "and causing the legitimate user's callback to fail. authId is in the Gate UI redirect URL. RISK REGISTRY."),
    ]:
        add_row(t12, t_data)

    # Apply label spacing to I-01/I-02 template sub-labels
    _i01i02_labels = {"Description", "Assets Involved", "Data Flow",
                      "Access Control", "Security Considerations", "Threat Assessment",
                      "Data Flow Diagram"}
    _heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name in _heading_styles:
            continue
        if p.text.strip() in _i01i02_labels and p.paragraph_format.space_before is None:
            apply_label_spacing(p)

    # ═════════════════════════════════════════════════════════════════════
    # 13. Programmatic interactions I-03 through I-16
    # ═════════════════════════════════════════════════════════════════════
    # Save references to template tables BEFORE InteractionBuilder adds new
    # tables (which shifts all subsequent indices).
    t_review = doc.tables[13]   # Security Checklist
    t_vuln = doc.tables[14]     # Vulnerability Management
    t_privacy = doc.tables[15]  # Privacy

    review_idx = find_paragraph_index(doc, "Review Checklist")
    if review_idx < 0:
        review_idx = find_paragraph_index(doc, "Security Considerations")
    insert_point = doc.paragraphs[review_idx]._p
    ib = InteractionBuilder(doc, insert_point)

    # ── I-03: Token Request ──
    ib.build_interaction("I-03", "Token Request (Client \u2192 Token Endpoint)",
        "POST /oauth2/token. Only endpoint wrapped with ClientAuthMiddleware.\n\n"
        "Processing flow:\n"
        "1. ClientAuth Middleware authenticates client (Basic/Post/None) with constant-time SHA-256.\n"
        "2. Validates grant_type presence, valid enum, and supported by grant handler provider.\n"
        "3. Validates grant_type is allowed for the specific app (per-app grant whitelist).\n"
        "4. Publishes TOKEN_ISSUANCE_STARTED observability event.\n"
        "5. Delegates to appropriate grant handler: ValidateGrant() then HandleGrant().\n"
        "6. Scope validation via scopeValidator.ValidateScopes() (currently pass-through).\n"
        "7. Issues refresh token if grant is authorization_code AND app allows refresh_token grant.\n"
        "8. Publishes TOKEN_ISSUED or TOKEN_ISSUANCE_FAILED event with duration_ms.\n\n"
        "Response headers: Cache-Control: no-store, Pragma: no-cache.\n"
        "Error mapping: server_error=500, unauthorized_client=401, all others=400.\n"
        "WWW-Authenticate: Basic returned on malformed Basic auth header.",
        [("OAuth Client","Client Auth MW","Token Endpoint \u2192 Grant Handler \u2192 Token Service")],
        "See DF-1 steps 11-13, DF-2, DF-3, DF-4.",
        "client_secret_basic/post/none per-app config. Dual credential submission (header + body) rejected "
        "with invalid_request. Client must be authorized for both grant_type and auth method.",
        [("Data Confidentiality","High [C-High]","Client creds + tokens"),
         ("Communication Medium","Network [M-NT]","HTTPS POST"),
         ("Transport Security","TLS 1.3",""),
         ("Authentication","Client Secret or None","Per-app config"),
         ("Accessibility","Public","")],
        [("1","Spoofing","Client Impersonation.","No","Constant-time SHA-256. PKCE for public. Dual cred rejected."),
         ("2","Replay","Auth Code Replay.","Yes",
          "RISK: SELECT then UPDATE (non-atomic). Two concurrent requests can both read ACTIVE code before either deactivates. Revocation of tokens from replayed codes not implemented (TODO in codebase). RISK REGISTRY."),
         ("3","Spoofing","redirect_uri mismatch at token.","No","Exact match against stored value."),
         ("4","Spoofing","PKCE verifier brute force.","No","43-128 chars, SHA-256 hashed. Infeasible."),
         ("5","Info Disc.","Token response caching.","No","Cache-Control: no-store + Pragma: no-cache."),
         ("6","Elevation","Scope escalation on refresh.","No","Scope intersection \u2014 cannot escalate."),
         ("7","DoS","Token endpoint abuse / brute-force.","Yes",
          "RISK: No rate limiting. Recommend implementing. RISK REGISTRY."),
         ("8","Info Disc.","Client secret in logs.","No","MaskString used. No cleartext secrets logged."),
         ("9","Spoofing","resource param mismatch.","No","Must match auth-time stored value."),
         ("10","Info Disc.","Timing attack on client auth.","No",
          "crypto/subtle.ConstantTimeCompare on SHA-256 hashes."),
         ("11","Spoofing","Unauthorized grant type.","No","Per-app grant whitelist enforced."),
         ("12","Info Disc.","Client secret stored in request context after authentication.","Partial",
          "OAuthClientInfo stores plain-text ClientSecret in request context (clientauth/context.go). "
          "Copied into tokenRequest. Increases risk from memory dumps or debug logging."),
         ("13","DoS","No request body size limits.","Yes",
          "RISK: No http.MaxBytesReader on token endpoint. Unbounded POST bodies accepted. "
          "Recommend adding body size limits. RISK REGISTRY."),
         ("14","Elevation","Client Credentials: no audience restriction enforcement.","Yes",
          "RISK: DetermineAudience uses resource param or falls back to client_id. No validation that client "
          "is authorized for the requested audience/resource. Any CC client can get tokens for any audience. "
          "Combined with no-op scope validation, a client can get tokens for any audience with any scope. RISK REGISTRY."),
         ("15","Spoofing","PKCE verification not constant-time.","Partial",
          "validatePlainChallenge and validateS256Challenge use != (regular string comparison) instead of "
          "subtle.ConstantTimeCompare. Enables timing side-channel to infer code_verifier/challenge characters. "
          "Practical impact limited for S256 (hash comparison) but plain method is fully vulnerable."),
         ("16","Tampering","No Content-Type validation on token endpoint.","Partial",
          "r.ParseForm() accepts both application/x-www-form-urlencoded and multipart/form-data. "
          "RFC 6749 Section 4.3.2 requires Content-Type: application/x-www-form-urlencoded. "
          "Accepting multipart could enable unexpected behavior with encoding differences.")])

    # ── I-04: Introspection ──
    ib.build_interaction("I-04", "Token Introspection (RS \u2192 Introspection)",
        "POST /oauth2/introspect with form parameter 'token' (required) and optional 'token_type_hint'.\n\n"
        "Processing flow:\n"
        "1. Parses form data; validates token parameter presence.\n"
        "2. Validates JWT via jwtService.VerifyJWT() (signature + expiry).\n"
        "3. If invalid: returns {active: false} (per RFC 7662 \u2014 no error details leaked).\n"
        "4. If valid: decodes JWT via jwt.DecodeJWT() and extracts claims.\n"
        "5. Returns: scope, client_id, username, exp, iat, nbf, sub, aud, iss, jti, token_type=Bearer.\n\n"
        "CRITICAL: No ClientAuthMiddleware applied. Any network caller can introspect any token.\n"
        "token_type_hint is accepted in request but ignored in processing.\n"
        "No token revocation checking (TODO in codebase).",
        [("Resource Server / Any","\u2014","Introspection \u2192 JWT Service (directly)")],
        "POST token. Thunder verifies JWT signature/expiry via jwtService, decodes and returns claims. "
        "Does NOT use TokenValidator \u2014 calls JWT Service directly.",
        "NO authentication or authorization. Any network caller can introspect any token. "
        "RFC 7662 Section 2.1 requires protected resource authentication.",
        [("Data Confidentiality","Medium [C-Medium]","Token metadata + claims returned"),
         ("Communication Medium","Network [M-NT]","HTTPS POST"),
         ("Transport Security","TLS 1.3",""),
         ("Authentication","None (GAP)","No client auth MW"),
         ("Accessibility","Public","Open to any caller")],
        [("1","Info Disc.","Unauthenticated introspection (RFC 7662 Section 2.1).","Yes",
          "RISK: Any caller can introspect any token. Require client auth. RISK REGISTRY."),
         ("2","Spoofing","Token fishing: probe with guessed tokens.","No",
          "JWTs high entropy. But endpoint auth would eliminate this."),
         ("3","Tampering","Missing revocation check.","Yes",
          "RISK: /oauth2/revoke not implemented. Cannot check revocation. RISK REGISTRY."),
         ("4","Info Disc.","token_type_hint accepted but ignored.","No",
          "No security impact. Hint only per RFC 7662."),
         ("5","Info Disc.","Missing Cache-Control: no-store on introspection response.","Yes",
          "RISK: Token metadata may be cached by proxies/browsers. Add Cache-Control: no-store.")])

    # ── I-05: UserInfo ──
    ib.build_interaction("I-05", "UserInfo (Client \u2192 UserInfo Endpoint)",
        "GET|POST /oauth2/userinfo. Bearer token required (openid scope).\n\n"
        "Processing flow:\n"
        "1. Extracts Bearer token from Authorization header.\n"
        "2. Verifies JWT via jwtService.VerifyJWT() (signature + expiry).\n"
        "3. Decodes JWT and extracts: sub, scope, client_id, grant_type, claims_req, claims_locales.\n"
        "4. Rejects client_credentials grant tokens (grant_type claim check).\n"
        "5. Retrieves OAuth app config via ApplicationService.GetOAuthApplication(client_id).\n"
        "6. Fetches LIVE user data from User Service: base attributes, groups (max 20 via Group Service), "
        "OU info (ouId, ouName, ouHandle via OU Service), userType.\n"
        "7. Builds claims based on: OIDC scope-to-claims mapping, per-app user_info.user_attributes filter, "
        "per-app scope_claims overrides, and claims_req from access token.\n"
        "8. Returns JSON with sub + filtered claims. Sets Cache-Control: no-store.\n\n"
        "Error responses: 401 for invalid/expired token, 403 for missing openid scope or CC token.",
        [("OAuth Client","\u2014","UserInfo \u2192 User Service + OU Service \u2192 User DB")],
        "Validates JWT, extracts sub/scope/client_id/grant_type/claims_req, fetches live data from "
        "User Service and OU Service, filters by scope + per-app user_info.user_attributes config.",
        "Bearer token signature + expiry verified. client_credentials rejected (grant_type claim). "
        "openid scope required. Per-app user_attributes filtering limits data exposure.",
        [("Data Confidentiality","High [C-High]","Returns PII; live DB data"),
         ("Communication Medium","Network [M-NT]","HTTPS"),
         ("Transport Security","TLS 1.3",""),
         ("Authentication","Bearer JWT","Signature + expiry validated"),
         ("Accessibility","Public","Requires valid token")],
        [("1","Info Disc.","Stolen Bearer token \u2192 PII (RFC 9700 Section 4.10).","Partial",
          "Bearer tokens are inherently stealable if intercepted. TLS 1.3 + short-lived JWTs mitigate but do not eliminate. "
          "Consider DPoP/mTLS per RFC 9700 Section 2.2.1. RISK REGISTRY (see RR-06)."),
         ("2","Info Disc.","Excessive data exposure.","No",
          "Filtered by scopes + per-app user_info.user_attributes."),
         ("3","Elevation","CC token at UserInfo.","No",
          "Explicitly rejected \u2014 checks grant_type claim."),
         ("4","Info Disc.","Stale data: UserInfo differs from ID token.","No",
          "By design: fetches live data. Clients should be aware."),
         ("5","Info Disc.","Group membership leak.","No",
          "Only if configured in per-app user_attributes. \u226420 groups."),
         ("6","Info Disc.","claims_req parameter from access token expands response.","No",
          "Claims still filtered by per-app user_attributes. Cannot request beyond app config."),
         ("7","Info Disc.","Missing security response headers.","Partial",
          "No HSTS, X-Content-Type-Options, or CSP headers. Cache-Control: no-store is set. "
          "Missing headers reduce defense-in-depth.")])

    # ── I-06: JWKS ──
    ib.build_interaction("I-06", "JWKS (External \u2192 JWKS Endpoint)",
        "GET /oauth2/jwks. Returns JWK Set with all public signing keys.\n\n"
        "Processing flow:\n"
        "1. Retrieves all certificates from CertificateService.GetCertificates().\n"
        "2. Iterates over each certificate, determining key type (RSA, EC, OKP).\n"
        "3. For each key, builds JWK with: kid (SHA-256 thumbprint), kty, use=sig, alg.\n"
        "4. Adds x5c chain (base64 DER), x5t (SHA-1 thumbprint, base64url), "
        "x5t#S256 (SHA-256 thumbprint, base64url).\n"
        "5. RSA keys include n, e. EC keys include crv (P-256/P-384/P-521), x, y. "
        "OKP keys include crv=Ed25519, x.\n"
        "6. Returns JSON {keys: [...]} array.\n\n"
        "NOTE: No caching headers set. Endpoint is regenerated on each request. "
        "kid is computed from certificate thumbprint (SHA-256), not configured. "
        "x5t uses SHA-1 (deprecated but still widely used for compatibility).",
        [("RS / Client","\u2014","JWKS \u2192 Certificate Service \u2192 PKI")],
        "GET request returns JWK Set with public keys. Keys loaded from PKI on each request.",
        "No auth (public). Keys are public by design. Private keys never exposed.",
        [("Data Confidentiality","Public","Public keys only"),
         ("Communication Medium","Network [M-NT]","HTTPS"),
         ("Transport Security","TLS 1.3","Prevents key injection"),
         ("Authentication","None","Public endpoint"),
         ("Accessibility","Public","")],
        [("1","Tampering","JWKS MITM: inject forged key.","No",
          "TLS 1.3. x5c chain + x5t#S256 for verification."),
         ("2","Info Disc.","Algorithm enumeration.","No",
          "Public. Server enforces specific algs."),
         ("3","Spoofing","Algorithm confusion (multi-key-type).","No",
          "JWT kid matches specific key. No alg=none."),
         ("4","Info Disc.","SHA-1 in x5t thumbprint.","No",
          "SHA-1 used for compatibility only. x5t#S256 also provided. "
          "SHA-256 thumbprint is primary kid."),
         ("5","DoS","No caching: repeated JWKS fetches hit PKI each time.","No",
          "Lightweight operation. Consider adding Cache-Control header for high-traffic deployments.")])

    # ── I-07: Discovery ──
    ib.build_interaction("I-07", "Discovery (External \u2192 /.well-known/*)",
        "Two separate endpoints serving server metadata:\n\n"
        "1. GET /.well-known/openid-configuration \u2014 OIDC Discovery 1.0 metadata.\n"
        "2. GET /.well-known/oauth-authorization-server \u2014 RFC 8414 AS metadata.\n\n"
        "Both endpoints build metadata dynamically from server config:\n"
        "\u2022 issuer, authorization_endpoint, token_endpoint, userinfo_endpoint, jwks_uri\n"
        "\u2022 introspection_endpoint, registration_endpoint\n"
        "\u2022 response_types_supported: [code]\n"
        "\u2022 grant_types_supported: [authorization_code, client_credentials, refresh_token, "
        "urn:ietf:params:oauth:grant-type:token-exchange]\n"
        "\u2022 token_endpoint_auth_methods_supported: [client_secret_basic, client_secret_post, none]\n"
        "\u2022 code_challenge_methods_supported: [S256, plain]\n"
        "\u2022 scopes_supported (OIDC + non-OIDC from ScopeService)\n"
        "\u2022 claims_supported, claims_parameter_supported: true\n"
        "\u2022 subject_types_supported: [public]\n\n"
        "Constants for /oauth2/revoke and /oauth2/logout defined but NOT registered as routes.\n"
        "NOTE: userinfo_endpoint field exists in the metadata model (omitempty) but is NOT populated \n"
        "in current responses. Clients must discover userinfo via token endpoint or documentation.",
        [("Client / RS","\u2014","Discovery Endpoints")],
        "GET returns JSON metadata document. Built dynamically from server runtime config.",
        "No auth (public per RFC 8414 / OIDC Discovery).",
        [("Data Confidentiality","Public","Capabilities + URLs"),
         ("Communication Medium","Network [M-NT]","HTTPS"),
         ("Transport Security","TLS 1.3",""),
         ("Authentication","None","Public"),
         ("Accessibility","Public","")],
        [("1","Tampering","Metadata MITM: redirect clients to malicious endpoints.","No",
          "TLS 1.3. RFC 8414 metadata (RFC 9700 Section 2.6)."),
         ("2","Info Disc.","Feature enumeration.","No","Public by design."),
         ("3","Tampering","Metadata mismatch between two endpoints.","No",
          "Same builder function produces both. Consistent."),
         ("4","Info Disc.","Advertised endpoints not implemented (revoke, logout).","Partial",
          "Constants defined but routes not registered. Endpoints return 404. Could mislead clients into believing revocation is available when it is not."),
         ("5","Info Disc.","Metadata advertises code_challenge_methods_supported: [S256, plain].","Yes",
          "RISK: Advertising plain method in metadata encourages clients to use it. RFC 9700 Section 2.1.1 recommends S256 only. "
          "Clients reading metadata may choose plain over S256. RISK REGISTRY (see RR-14).")])

    # ── I-08: DCR ──
    ib.build_interaction("I-08", "DCR (External \u2192 DCR Endpoint)",
        "POST /oauth2/dcr/register. RFC 7591 Dynamic Client Registration.\n\n"
        "Processing flow:\n"
        "1. If dcr.insecure=false: checks security.GetPermissions for system-level permission.\n"
        "2. Parses JSON request body with client metadata.\n"
        "3. client_id auto-generated (UUID) \u2014 cannot be client-specified.\n"
        "4. Public client auto-detection: if token_endpoint_auth_method=none, marks as public client.\n"
        "5. Public clients auto-enforce: pkce_required=true. client_credentials grant removal\n"
        "   handled downstream by Application Service during app creation (not by DCR directly).\n"
        "6. Validates: redirect_uris (required for auth_code), grant_types vs response_types consistency, "
        "auth_method supported, jwks/jwks_uri mutual exclusivity.\n"
        "7. Maps jwks_uri/jwks to application certificates.\n"
        "8. Generates client_secret (random), SHA-256 hashed before storage.\n"
        "9. Creates application via ApplicationService.\n"
        "10. Returns: client_id, client_secret, client_secret_expires_at=0 (never), "
        "plus echoed metadata. Additional fields: client_uri, logo_uri, tos_uri, policy_uri, contacts.",
        [("Client / Dev","Permission Check","DCR \u2192 Application Service")],
        "POST metadata \u2192 validate \u2192 auto-detect public client \u2192 create app \u2192 return client_id + client_secret.",
        "dcr.insecure=false \u2192 system permission (security.GetPermissions). "
        "Validates: redirect_uris, grant_types, auth_method, response_types, jwks exclusivity. "
        "Public clients auto-enforced (PKCE required, no CC grant).",
        [("Data Confidentiality","High [C-High]","Returns client_secret"),
         ("Communication Medium","Network [M-NT]","HTTPS POST"),
         ("Transport Security","TLS 1.3","Secret over TLS only"),
         ("Authentication","Configurable","insecure flag"),
         ("Accessibility","Configurable","")],
        [("1","Spoofing","Unauth registration when dcr.insecure=true.","Yes",
          "RISK: Open registration. Use insecure=false in prod. Initial access tokens per RFC 7591 Section 3. RISK REGISTRY."),
         ("2","DoS","Registration flooding.","Yes",
          "RISK: No rate limiting. RISK REGISTRY."),
         ("3","Elevation","Excessive grants/scopes via DCR.","No",
          "Validated against supported capabilities. Public clients auto-enforce PKCE."),
         ("4","Info Disc.","client_secret in response; never expires.","No",
          "Returned once over TLS. client_secret_expires_at=0 \u2014 recommend rotation support."),
         ("5","Tampering","JWKS/JWKS_URI injection.","No",
          "Mutually exclusive validation. Mapped to app certificates."),
         ("6","Elevation","Public client requests client_credentials grant.","No",
          "Auto-removed during public client detection. Cannot bypass."),
         ("7","Spoofing","Client-specified client_id.","No",
          "client_id is server-generated UUID. Client input ignored."),
         ("8","DoS","No request body size limits.","Yes",
          "RISK: No http.MaxBytesReader. DCR accepts unbounded JSON bodies. "
          "Combined with insecure mode (no auth), attacker can send arbitrarily large payloads to exhaust memory/storage. RISK REGISTRY.")])

    # ── I-09: Token Exchange ──
    ib.build_interaction("I-09", "Token Exchange (Client \u2192 Token Endpoint, RFC 8693)",
        "grant_type=urn:ietf:params:oauth:grant-type:token-exchange.\n\n"
        "Processing flow:\n"
        "1. Validates required params: subject_token, subject_token_type.\n"
        "2. Verifies subject_token JWT (signature + expiry via jwtService).\n"
        "3. If subject_token contains 'assurance' claim, treats as auth assertion with stricter validation:\n"
        "   \u2022 audience must match configured audience OR app_id.\n"
        "   \u2022 authorized_permissions claim used for scope (not requested scope).\n"
        "4. Non-assertion tokens: audience priority = audience param > resource > subject_token aud > client_id.\n"
        "5. If actor_token provided: verifies JWT, nests act claim chain in issued token.\n"
        "6. Scope subsetting: requested scope intersected with subject token scope.\n"
        "7. Builds new access token via tokenService with delegation act claim.\n"
        "8. Returns: access_token, issued_token_type=urn:ietf:params:oauth:token-type:access_token, "
        "token_type=Bearer, expires_in, scope.\n\n"
        "NOTE: External JWKS validation for subject_token not fully implemented (TODO in codebase). "
        "Currently only validates tokens signed by this server's keys.",
        [("OAuth Client","Client Auth MW","Token Exchange Handler \u2192 Token Service")],
        "See DF-4: Token Exchange flow. Validates subject/actor JWTs, computes audience priority, "
        "detects auth assertions (assurance claim), subsets scope, builds new token with act claim.",
        "Client auth required. subject/actor tokens validated (sig, expiry, aud). "
        "Auth assertions: audience must match config or app_id. Scope cannot escalate.",
        [("Data Confidentiality","High [C-High]","Existing + new tokens"),
         ("Communication Medium","Network [M-NT]","HTTPS POST"),
         ("Transport Security","TLS 1.3",""),
         ("Authentication","Client Secret","Required"),
         ("Accessibility","Public","Requires valid creds")],
        [("1","Elevation","Escalate via exchange.","No",
          "Inherits subject. Scope equal or narrower. act chain tracked."),
         ("2","Spoofing","Stolen subject_token reuse.","No",
          "Client auth + token signature/expiry validation."),
         ("3","Tampering","Audience manipulation.","No",
          "Clear priority order. Strict for auth assertions."),
         ("4","Spoofing","Forged actor_token.","No",
          "JWT signature + expiry checked. act claim nests chain."),
         ("5","Spoofing","Auth assertion abuse (assurance claim).","No",
          "Stricter audience. authorized_permissions for scope."),
         ("6","Elevation","External token accepted without JWKS validation.","Yes",
          "RISK: Currently only validates tokens signed by this server. External JWKS TODO. RISK REGISTRY."),
         ("7","Elevation","Scope escalation via requested_scope param.","No",
          "Scope intersected with subject token. Cannot request more than subject has."),
         ("8","Spoofing","Subject token decoded before signature verification.","Partial",
          "ValidateSubjectToken first decodes JWT payload to extract issuer for key routing, then verifies signature. "
          "Issuer routing decision based on unverified claims. Currently mitigated by accepting only internal issuers, "
          "but if external issuers are added (TODO exists), this pattern becomes dangerous."),
         ("9","Info Disc.","Error messages leak internal validation details.","Yes",
          "RISK: Error response includes fmt.Sprintf('Invalid subject_token: %s', err.Error()) which leaks internal "
          "error messages about token structure, missing claims, signature failures, and issuer validation. RISK REGISTRY.")],
        diagram_path=os.path.join(DIAGRAMS_DIR, "flow-token-exchange.png"))

    # ── I-10: Database Ops ──
    ib.build_interaction("I-10", "Database Operations (OAuth \u2192 DB)",
        "Internal data access for OAuth state management.\n\n"
        "Tables:\n"
        "\u2022 AUTHORIZATION_REQUEST: Stores auth request context (redirect_uri, scope, state, PKCE, "
        "resource, claims, claims_locales, oidc_scopes, non_oidc_scopes, response_type, "
        "user_attributes as JSON). 10-min TTL, single-use (deleted on read).\n"
        "\u2022 AUTHORIZATION_CODE: Stores auth codes with all grant data (PKCE challenge/method, "
        "redirect_uri, scope, state, resource, claims, user_attributes as JSON). "
        "State machine: ACTIVE \u2192 INACTIVE on first use.\n\n"
        "IMPORTANT: Auth code consumption (SELECT then UPDATE) and auth request consumption \n"
        "(SELECT then DELETE) are non-atomic \u2014 no transaction wrapping. This creates a small \n"
        "race condition window for concurrent requests with the same code/authId.\n\n"
        "NOTE: Auth code SELECT query does NOT filter by STATE=ACTIVE in SQL; state validation \n"
        "happens in the grant handler in-memory after retrieval.\n\n"
        "All queries use parameterized SQL via DBQuery with unique IDs for traceability. "
        "DEPLOYMENT_ID included in all WHERE clauses as composite PK for multi-tenant isolation. "
        "Complex data (user_attributes, claims) stored as JSON blobs.",
        [("OAuth Module","DBClient","Runtime DB / Thunder DB")],
        "Internal CRUD operations. Auth request: INSERT (create) \u2192 SELECT then DELETE (non-atomic). "
        "Auth code: INSERT (create) \u2192 SELECT then UPDATE status (non-atomic).",
        "In-process for SQLite. Authenticated for PostgreSQL. DEPLOYMENT_ID isolation. "
        "JSON blob storage for complex nested data structures.",
        [("Data Confidentiality","High [C-High]","Auth codes, contexts, configs"),
         ("Communication Medium","Internal [M-IN]",""),
         ("Transport Security","N/A or DB TLS",""),
         ("Authentication","DB-level",""),
         ("Accessibility","Internal","")],
        [("1","Tampering","SQL Injection.","No","Parameterized queries via DBQuery. Never concatenated."),
         ("2","Info Disc.","Code theft from DB.","No",
          "Single-use, short-lived. DEPLOYMENT_ID isolation. State machine prevents reuse."),
         ("3","Tampering","Cross-tenant access.","No",
          "DEPLOYMENT_ID in all WHERE clauses as composite PK."),
         ("4","Tampering","JSON blob manipulation in DB.","No",
          "JSON blobs are server-controlled. No client input directly to JSON columns."),
         ("5","Info Disc.","Auth context data residue.","No",
          "Contexts deleted on read. Codes set INACTIVE. Consider periodic cleanup job."),
         ("6","Replay","Auth code race condition: concurrent token requests with same code.","Yes",
          "RISK: SELECT and UPDATE are non-atomic. Two concurrent requests may both read ACTIVE code "
          "before either deactivates it. Use SELECT...FOR UPDATE or wrap in transaction. RISK REGISTRY."),
         ("7","Replay","Auth request context race condition.","No",
          "Low risk: authId is UUIDv7 with 10-min TTL. Attacker needs both authId and valid assertion."),
         ("8","DoS","No expired auth request/code DB cleanup.","Yes",
          "RISK: No cleanup, cron, scheduled, purge, or garbage collection mechanism found anywhere in the OAuth module. "
          "Expired auth requests and deactivated auth codes accumulate indefinitely. Tables grow unboundedly, "
          "degrading query performance and consuming storage. RISK REGISTRY."),
         ("9","Replay","Auth code not checked for expiry before deactivation.","Yes",
          "RISK: Auth code SELECT query does not filter by expiry in SQL. Expired codes can be retrieved and state checked in-memory. "
          "Combined with race condition (T6), an expired code could potentially be used if deactivation is not atomic. RISK REGISTRY.")])

    # ── I-11: Observability Events ──
    ib.build_interaction("I-11", "Observability Events (Token Service \u2192 Events)",
        "In-process event publishing during token requests only.\n\n"
        "Events published:\n"
        "\u2022 TOKEN_ISSUANCE_STARTED: At beginning of token request processing.\n"
        "\u2022 TOKEN_ISSUED: On successful token issuance.\n"
        "\u2022 TOKEN_ISSUANCE_FAILED: On failure, with error details.\n\n"
        "Event data includes: client_id, grant_type, scope, error, error_code, "
        "error_type (client_error vs server_error), duration_ms.\n\n"
        "NOTE: Events only published for /oauth2/token endpoint. No events for: "
        "introspection, userinfo, JWKS, discovery, DCR, or authorization. "
        "error_type distinguishes client_error (4xx) from server_error (5xx) for operational analysis.",
        [("Token Service","Event Publisher","Internal Event System")],
        "In-process event publishing via EventPublisher. Events fired at start, success, and failure "
        "of token processing. Contains operational data only.",
        "Internal mechanism. No external API. Events consumed by observability infrastructure.",
        [("Data Confidentiality","Medium","Operational data"),
         ("Communication Medium","Internal [M-IN]",""),
         ("Transport Security","N/A",""),
         ("Authentication","N/A",""),
         ("Accessibility","Internal","")],
        [("1","Info Disc.","Sensitive data in events if exported.","No",
          "Operational data only. No secrets/tokens/PII. client_id and scope are non-sensitive."),
         ("2","Tampering","Event flooding on high volume.","No",
          "Lightweight metadata. Proportional to requests."),
         ("3","Info Disc.","Missing event coverage for non-token endpoints.","No",
          "By design for now. Consider adding events for DCR, introspection for audit trail.")])

    # ── I-12: Scope Validation ──
    ib.build_interaction("I-12", "Scope Validation (OAuth \u2192 Scope Validator)",
        "Scope handling spans multiple components with different enforcement levels.\n\n"
        "Scope processing pipeline:\n"
        "1. Authorization endpoint: Separates OIDC scopes (openid, profile, email, phone, address) "
        "from non-OIDC scopes. Computes required user attributes from OIDC scope mappings.\n"
        "2. Auth callback: Overwrites non-OIDC scopes with authorized_permissions from Flow assertion. "
        "OIDC scopes preserved from original request.\n"
        "3. Token endpoint: scopeValidator.ValidateScopes() called \u2014 CURRENTLY RETURNS ALL REQUESTED "
        "SCOPES (no-op). No per-app scope whitelist enforcement at issuance.\n"
        "4. Refresh token grant: Scope intersection \u2014 requested scope must be subset of original. "
        "This is the ONLY point where scope downscoping is enforced.\n"
        "5. Claim building: Per-app scope_claims configuration maps scopes to specific claims. "
        "This is where scope-to-claim translation actually happens and can override OIDC standard mappings.\n\n"
        "CRITICAL: Token carries full requested scope without per-app whitelist. "
        "Enforcement only at claim building (what data is in the token) and refresh (cannot escalate).",
        [("OAuth Module","Scope Validator","Token Service / Claim Builder")],
        "Called during token processing. ValidateScopes returns all requested scopes unconditionally. "
        "Real enforcement: refresh downscoping (intersection), claim building (per-app scope_claims).",
        "Currently returns all requested scopes without per-app policy. "
        "Meaningful enforcement only at claim-building stage and refresh token intersection.",
        [("Data Confidentiality","Medium","Scopes determine resource/claim access"),
         ("Communication Medium","Internal [M-IN]",""),
         ("Transport Security","N/A",""),
         ("Authentication","N/A",""),
         ("Accessibility","Internal","")],
        [("1","Elevation","Scope bypass: arbitrary scopes granted.","Yes",
          "RISK: No per-app scope whitelist at token issuance. Token carries full requested scope. "
          "Enforcement only at claim building. RISK REGISTRY."),
         ("2","Elevation","scope_claims override OIDC mappings.","No",
          "Admin-configured per-app. Cannot be influenced by clients."),
         ("3","Elevation","Non-OIDC scope expansion.","No",
          "Non-OIDC scopes overwritten by authorized_permissions from assertion, not merged."),
         ("4","Info Disc.","Scope leaks in token claims.","No",
          "scope claim in token is visible to client. By design per OAuth 2.0 spec.")])

    # ── I-13: Client Authentication Middleware ──
    ib.build_interaction("I-13", "Client Authentication Middleware (Token Endpoint Guard)",
        "Middleware applied exclusively to POST /oauth2/token. Authenticates the OAuth client "
        "before any grant handler executes.\n\n"
        "Processing flow:\n"
        "1. Checks for Authorization: Basic header \u2192 decodes base64, extracts client_id:client_secret.\n"
        "2. If no Basic header, checks POST body for client_id + client_secret params.\n"
        "3. DUAL CREDENTIAL REJECTION: If credentials found in BOTH header AND body, returns "
        "invalid_request/400 (per RFC 6749 Section 2.3).\n"
        "4. If no credentials at all, checks POST body for client_id only (public client / 'none' auth).\n"
        "5. Retrieves OAuth app config via ApplicationService.GetOAuthApplication(client_id).\n"
        "6. Validates auth method matches app config (client_secret_basic, client_secret_post, none).\n"
        "7. For secret-based auth: SHA-256 hash of provided secret, constant-time comparison "
        "(crypto/subtle.ConstantTimeCompare) against stored hash.\n"
        "8. On success: sets authenticated client in request context for downstream handlers.\n\n"
        "Error responses:\n"
        "\u2022 invalid_client (401) + WWW-Authenticate: Basic \u2014 for bad credentials.\n"
        "\u2022 invalid_request (400) \u2014 for dual credential or missing client_id.\n"
        "\u2022 unauthorized_client (401) \u2014 for wrong auth method.",
        [("OAuth Client","\u2014","Client Auth MW \u2192 Application Service")],
        "Extracts credentials from header or body. Validates against app config. "
        "Sets authenticated client in context.",
        "client_secret_basic: Basic header with base64(client_id:secret). "
        "client_secret_post: client_id + client_secret in POST body. "
        "none: client_id only, no secret (public clients). "
        "Constant-time SHA-256 comparison prevents timing attacks.",
        [("Data Confidentiality","High [C-High]","Client credentials in transit"),
         ("Communication Medium","Network [M-NT]","HTTPS POST"),
         ("Transport Security","TLS 1.3","Credentials over TLS"),
         ("Authentication","Client Secret","Per-app method"),
         ("Accessibility","Public","Requires valid client_id")],
        [("1","Spoofing","Timing attack on secret comparison.","No",
          "crypto/subtle.ConstantTimeCompare on SHA-256 hashes. No timing leakage."),
         ("2","Spoofing","Dual credential injection (header + body).","No",
          "Explicitly rejected with invalid_request/400 per RFC 6749 Section 2.3."),
         ("3","Spoofing","Brute-force client secret.","Yes",
          "RISK: No rate limiting or account lockout. RISK REGISTRY."),
         ("4","Info Disc.","Client secret in Basic header logged.","No",
          "MaskString applied. No cleartext credentials in logs."),
         ("5","Spoofing","Auth method mismatch.","No",
          "Per-app config enforced. Cannot use basic if app configured for post."),
         ("6","Spoofing","Public client impersonation.","No",
          "PKCE enforced for public clients. client_id alone insufficient.")])

    # ── I-14: PKCE Lifecycle ──
    ib.build_interaction("I-14", "PKCE Lifecycle (Cross-Endpoint Security Binding)",
        "PKCE (RFC 7636) spans authorization and token endpoints.\n\n"
        "Authorization endpoint (challenge):\n"
        "1. code_challenge and code_challenge_method extracted from query params.\n"
        "2. For public clients: code_challenge is REQUIRED. Rejection if missing.\n"
        "3. Supported methods: S256 (SHA-256 + base64url), plain.\n"
        "4. code_challenge length: 43 chars for S256, 43-128 chars for plain (per RFC 7636 Section 4.2).\n"
        "5. Stored with authorization request context and persisted to AUTHORIZATION_CODE record.\n\n"
        "Token endpoint (verification):\n"
        "1. code_verifier extracted from token request params.\n"
        "2. If stored code_challenge exists: code_verifier REQUIRED.\n"
        "3. If NO stored code_challenge: code_verifier REJECTED (PKCE downgrade prevention per "
        "RFC 9700 Section 4.8.2).\n"
        "4. Computes challenge from verifier using stored method (S256 or plain).\n"
        "5. Constant-time comparison of computed vs stored challenge.\n\n"
        "Security binding: Links authorization request to token request via cryptographic proof. "
        "Prevents code interception attacks even without client secret.",
        [("OAuth Client","\u2014","Auth Endpoint \u2192 (stored) \u2192 Token Endpoint")],
        "Challenge stored at authorization time. Verifier validated at token time. "
        "Cross-endpoint binding via AUTHORIZATION_CODE table.",
        "Required for public clients. Optional for confidential. "
        "S256 recommended (SHA-256 hash). Downgrade prevention enforced.",
        [("Data Confidentiality","Medium","code_verifier is secret"),
         ("Communication Medium","Network [M-NT]","HTTPS"),
         ("Transport Security","TLS 1.3","Verifier over TLS"),
         ("Authentication","Proof-of-possession","Cryptographic binding"),
         ("Accessibility","Public","Part of standard flow")],
        [("1","Spoofing","Code interception without verifier.","No",
          "Verifier required to exchange code. Attacker has code but not verifier."),
         ("2","Tampering","PKCE downgrade: strip code_challenge.","No",
          "code_verifier rejected if no stored challenge (RFC 9700 Section 4.8.2)."),
         ("3","Spoofing","PKCE verifier brute force.","No",
          "43-128 chars, SHA-256 hashed. Computationally infeasible."),
         ("4","Tampering","Method downgrade: S256 \u2192 plain.","No",
          "Method stored with challenge. Must match at verification."),
         ("5","Replay","Reuse code_challenge across requests.","No",
          "Code is single-use. Challenge bound to specific code."),
         ("6","Spoofing","PKCE plain method allowed — security downgrade (RFC 9700 Section 2.1.1).","Yes",
          "RISK: plain method is explicitly supported and advertised in metadata. With plain, code_challenge == code_verifier "
          "(no hashing). An attacker who intercepts code_challenge can derive code_verifier trivially. "
          "RFC 9700 Section 2.1.1 and OAuth Security BCP recommend S256 only. Should disable plain or make S256 mandatory. RISK REGISTRY.")])

    # ── I-15: Token Building and Signing ──
    ib.build_interaction("I-15", "Token Building and Signing (Internal \u2192 JWT Service)",
        "Internal component that constructs and signs JWT tokens.\n\n"
        "Token types built:\n"
        "\u2022 Access Token: sub, aud, iss, scope, client_id, grant_type, claims_req, claims_locales, "
        "user attributes (per-app access_token.user_attributes), act (delegation chain), resource.\n"
        "\u2022 Refresh Token: access_token_sub, access_token_aud, scope, client_id, grant_type, "
        "metadata for re-derivation of access tokens.\n"
        "\u2022 ID Token: sub, iss, aud=client_id, auth_time, OIDC standard claims filtered by scope, "
        "per-app id_token.user_attributes, at_hash.\n\n"
        "Signing process:\n"
        "1. Token claims assembled by ClaimBuilder based on grant type and per-app config.\n"
        "2. Per-app scope_claims can override standard OIDC scope-to-claims mappings.\n"
        "3. User attributes filtered by per-app user_attributes configuration.\n"
        "4. JWT signed via jwtService.GenerateJWT() using preferred_key_id from config.\n"
        "5. Signing algorithms: RS256 (RSA), ES256/ES384/ES512 (ECDSA), EdDSA (Ed25519).\n"
        "6. Token validity configurable per-app (access_token.validity_period, etc.).\n\n"
        "Claims filtering pipeline: scope \u2192 OIDC mapping \u2192 scope_claims override \u2192 "
        "user_attributes filter \u2192 final claims.",
        [("Token Service","\u2014","Claim Builder \u2192 JWT Service \u2192 PKI")],
        "Claims assembled from grant data, per-app config, and user data. "
        "Signed by JWT Service using PKI key (preferred_key_id). Multiple algorithm support.",
        "Internal process. No direct external access. Keys stored on filesystem "
        "with restricted permissions.",
        [("Data Confidentiality","High [C-High]","Token claims contain user data"),
         ("Communication Medium","Internal [M-IN]",""),
         ("Transport Security","N/A",""),
         ("Authentication","N/A","Internal"),
         ("Accessibility","Internal","")],
        [("1","Tampering","Token forgery.","No",
          "Signed with private key. Key compromise = full token forgery."),
         ("2","Info Disc.","Excessive claims in tokens.","No",
          "Per-app user_attributes filtering. scope_claims configuration."),
         ("3","Info Disc.","Private key leakage.","No",
          "Filesystem with restricted permissions. Recommend Vault/KMS for production."),
         ("4","Tampering","Algorithm confusion.","No",
          "kid in JWT header binds to specific key and algorithm. No alg=none."),
         ("5","Info Disc.","scope_claims override leaks unexpected data.","No",
          "Admin-configured per-app. Cannot be influenced by clients or end users."),
         ("6","Elevation","Token validity too long.","No",
          "Configurable per-app. Default from server config. Recommend short-lived access tokens."),
         ("7","Replay","Missing OIDC nonce parameter support.","Yes",
          "RISK: nonce is not extracted from authorization request, not stored, and not included in ID tokens. "
          "Without nonce, ID token replay attacks are possible. Clients cannot correlate ID tokens to auth requests. "
          "Required by OIDC Core Section 3.1.2.1 for implicit flows, RECOMMENDED for code flow. RISK REGISTRY."),
         ("8","Spoofing","Missing at_hash and c_hash claims in ID token.","Yes",
          "RISK: at_hash and c_hash are completely absent from ID token builder. at_hash is REQUIRED by OIDC Core "
          "Section 3.3.2.11 for code flow when ID token is returned alongside access token. Without at_hash, clients "
          "cannot detect access token substitution attacks. RISK REGISTRY."),
         ("9","Info Disc.","Refresh token is JWT with sensitive claims.","Partial",
          "Refresh tokens contain access_token_sub (user ID), access_token_aud, grant_type, scope, and "
          "access_token_user_attributes. If logged, cached, or intercepted, exposes full permission set and user identity. "
          "Opaque refresh tokens would be safer."),
         ("10","Spoofing","Token substitution: no typ header or token_use claim.","Partial",
          "Access tokens and ID tokens are both JWTs signed with the same key. No explicit typ header (e.g. at+jwt) "
          "or token_use claim differentiates them. A resource server without proper claim validation could accept "
          "an access token where an ID token is expected, or vice versa.")])

    # ── I-16: CORS Middleware ──
    ib.build_interaction("I-16", "CORS Middleware (All OAuth Endpoints)",
        "CORS middleware applied to ALL OAuth endpoints via middleware.WithCORS.\n\n"
        "Configuration per endpoint:\n"
        "\u2022 Allowed origins: From server configuration.\n"
        "\u2022 Allowed methods: Per-endpoint (e.g., GET for /authorize, POST for /token).\n"
        "\u2022 Credentials: Access-Control-Allow-Credentials included.\n"
        "\u2022 Preflight: OPTIONS requests return 204 with CORS headers.\n\n"
        "COMPLIANCE ISSUE (RFC 9700 Section 2.6): CORS is applied to /oauth2/authorize. "
        "The authorization endpoint should NOT have CORS since it is a redirect-based endpoint "
        "accessed via browser navigation, not XHR/fetch. Enabling CORS on it could allow "
        "cross-origin scripted access to authorization responses.\n\n"
        "CORS on /oauth2/token, /oauth2/userinfo is standard for SPAs accessing these endpoints "
        "via fetch/XHR from allowed origins.",
        [("Browser / SPA","\u2014","CORS MW \u2192 OAuth Endpoints")],
        "Browser sends preflight OPTIONS; middleware returns CORS headers. "
        "Actual request includes Origin header; middleware validates against allowed origins.",
        "Origin validation against configured allowed origins. "
        "Credentials flag set. Per-endpoint method control.",
        [("Data Confidentiality","Low","CORS headers are informational"),
         ("Communication Medium","Network [M-NT]","HTTPS"),
         ("Transport Security","TLS 1.3",""),
         ("Authentication","None","CORS is pre-authentication"),
         ("Accessibility","Public","Browser security policy")],
        [("1","Tampering","CORS on authorization endpoint.","Yes",
          "RISK: RFC 9700 Section 2.6 violation. Should remove CORS from /oauth2/authorize. RISK REGISTRY."),
         ("2","Spoofing","Overly permissive allowed origins.","No",
          "Origins from server config. Recommend restrictive origin list."),
         ("3","Info Disc.","Credentials flag enables cookie/auth forwarding.","No",
          "Required for authenticated endpoints (token, userinfo). Expected behavior."),
         ("4","Spoofing","CORS bypass via non-browser client.","No",
          "CORS is browser-enforced. Server-side clients unaffected. "
          "Server still requires auth where applicable."),
         ("5","Info Disc.","Missing security response headers on ALL OAuth endpoints.","Yes",
          "RISK: No HSTS, X-Content-Type-Options, X-Frame-Options, or CSP headers set globally. "
          "Only token and userinfo set Cache-Control: no-store. Missing headers reduce defense-in-depth "
          "against XSS, MIME sniffing, clickjacking. RISK REGISTRY.")])

    # ═════════════════════════════════════════════════════════════════════
    # 14. Security Checklist (T13)
    # ═════════════════════════════════════════════════════════════════════
    print("  Filling checklists...")
    t13 = t_review
    set_table_cell_margins(t13)
    checklist = {
        0: ("Yes", "OAuth params validated. Redirect URIs exact-matched. resource validated as absolute URI."),
        1: ("No", "No rate limiting on any OAuth endpoint. Token + DCR highest risk."),
        2: ("Partial", "Token: Client Auth MW. UserInfo: Bearer. Introspection/DCR(insecure): no auth."),
        3: ("Yes", "Per-app grant types, auth methods. Public clients enforce PKCE."),
        4: ("N/A", "No defaults. Secrets generated at registration, SHA-256 hashed."),
        5: ("No", "Open-source. Security relies on implementation soundness."),
        6: ("Yes", "PR-based code reviews on GitHub."),
        7: ("Yes", "gosec, staticcheck."),
        8: ("Yes", "go.mod + go vuln for CVE scanning."),
        9: ("To Do", "Recommend DAST (OWASP ZAP) for OAuth endpoints."),
        10: ("Partial", "Logging + observability events. Recommend structured audit logs."),
        11: ("N/A", "Config at deploy time, not runtime."),
        12: ("Partial", "Health check + DB pools. Document RPO/RTO."),
        13: ("Yes", "TLS 1.3 transit. JWT signed (RS256+ECDSA+EdDSA). Encryption key at rest."),
        14: ("Partial", "Keys on filesystem. Recommend Vault/KMS."),
        15: ("Yes", "MaskString for sensitive data. No cleartext secrets/PII in logs."),
    }
    for row_idx, (state, comment) in checklist.items():
        data_row = row_idx + 1
        if data_row < len(t13.rows):
            set_cell_text(t13.rows[data_row].cells[1], state)
            set_cell_text(t13.rows[data_row].cells[2], comment)

    # ═════════════════════════════════════════════════════════════════════
    # 15. Vulnerability Management (T14)
    # ═════════════════════════════════════════════════════════════════════
    t14 = t_vuln
    set_table_cell_margins(t14)
    for i, txt in enumerate([
        "Go deps via go.mod. go vuln + dep scanning. Patches on release cycles.",
        "OS/container patches by deployment team. Docker images updated.",
        "No EOL components. Go latest stable, PostgreSQL (active), SQLite (active).",
    ]):
        if i+1 < len(t14.rows) and len(t14.rows[i+1].cells) >= 2:
            set_cell_text(t14.rows[i+1].cells[1], txt)

    # ═════════════════════════════════════════════════════════════════════
    # 16. Privacy (T15)
    # ═════════════════════════════════════════════════════════════════════
    t15 = t_privacy
    set_table_cell_margins(t15)
    privacy = {
        0: ("Yes","User auth, token claims, UserInfo. Consent via scope authorization."),
        1: ("Yes","Encryption at rest. JWTs signed + timed. Cache-Control: no-store everywhere."),
        2: ("N/A","Privacy policies by deploying org."),
        3: ("Yes","Filtered by scopes + per-app user_attributes + per-app scope_claims."),
        4: ("Yes","Codes: configurable. Contexts: 10-min. Tokens: per-app."),
        5: ("Partial","Codes/contexts auto-cleaned. Token revocation TODO. Account deletion should cascade."),
        6: ("To Do","Ensure OAuth data in WSO2 Data Inventory."),
    }
    for row_idx, (state, comment) in privacy.items():
        data_row = row_idx + 1
        if data_row < len(t15.rows):
            set_cell_text(t15.rows[data_row].cells[1], state)
            set_cell_text(t15.rows[data_row].cells[2], comment)

    # ═════════════════════════════════════════════════════════════════════
    # 17. Risk Registry
    # ═════════════════════════════════════════════════════════════════════
    print("  Writing risk registry...")
    idx_rr = find_paragraph_index(doc, "Risk registry entries")
    if idx_rr >= 0:
        for pi in range(idx_rr+1, min(idx_rr+5, len(doc.paragraphs))):
            p = doc.paragraphs[pi]
            if not p.text.strip() or p.text.strip().startswith(("<","[")):
                set_paragraph_text(p,
                    "RR-01: Introspection lacks auth (I-04 T1) \u2014 Critical \u2014 Open\n"
                    "  \u2192 Require client_secret_basic or IP restriction. Any caller can introspect any token.\n\n"
                    "RR-02: Token revocation not implemented (I-04 T3) \u2014 Critical \u2014 Open\n"
                    "  \u2192 Implement /oauth2/revoke + revocation checks. Stolen tokens irrevocable until expiry.\n\n"
                    "RR-03: DCR unauthenticated when insecure=true (I-08 T1) \u2014 High \u2014 Open\n"
                    "  \u2192 Always insecure=false in prod. Initial access tokens.\n\n"
                    "RR-04: No rate limiting on OAuth endpoints (I-03 T7, I-08 T2, I-13 T3) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Implement rate limiting on token, DCR, introspection.\n\n"
                    "RR-05: Scope validation no-op (I-12 T1) \u2014 High \u2014 Open\n"
                    "  \u2192 Implement per-app scope whitelist at token issuance. Currently returns all requested scopes.\n\n"
                    "RR-06: Bearer tokens not sender-constrained (I-05 T1) \u2014 Low \u2014 Open\n"
                    "  \u2192 Consider DPoP (RFC 9449) or mTLS (RFC 8705).\n\n"
                    "RR-07: client_secret never expires (I-08 T4) \u2014 Low \u2014 Accepted\n"
                    "  \u2192 Recommend secret rotation support.\n\n"
                    "RR-08: CORS on authorization endpoint (I-16 T1) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Remove CORS from /oauth2/authorize per RFC 9700 Section 2.6.\n\n"
                    "RR-09: External token JWKS validation missing (I-09 T6) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Implement external JWKS fetching for token exchange subject_tokens.\n\n"
                    "RR-10: Auth code consumption race condition (I-03 T2, I-10 T6) \u2014 High \u2014 Open\n"
                    "  \u2192 Non-atomic SELECT+UPDATE. Use SELECT...FOR UPDATE or transaction. Revoke tokens from replayed codes.\n\n"
                    "RR-11: Clickjacking on authorize endpoint (I-01 T15) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Add X-Frame-Options: DENY or CSP frame-ancestors 'none'.\n\n"
                    "RR-12: No Referrer-Policy header (I-01 T16) \u2014 Low \u2014 Open\n"
                    "  \u2192 Add Referrer-Policy: no-referrer on auth responses.\n\n"
                    "RR-13: Introspection missing Cache-Control (I-04 T5) \u2014 Low \u2014 Open\n"
                    "  \u2192 Add Cache-Control: no-store to introspection responses.\n\n"
                    "RR-14: PKCE plain method allowed (I-01 T18, I-14 T6, I-07 T5) \u2014 High \u2014 Open\n"
                    "  \u2192 Disable plain method. Make S256 mandatory per RFC 9700 Section 2.1.1.\n\n"
                    "RR-15: Missing OIDC nonce parameter support (I-15 T7) \u2014 High \u2014 Open\n"
                    "  \u2192 Extract nonce from auth request, store, include in ID token per OIDC Core Section 3.1.2.1.\n\n"
                    "RR-16: No request body size limits on any endpoint (I-03 T13, I-08 T8) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Add http.MaxBytesReader to all POST endpoints. Recommend 1MB limit.\n\n"
                    "RR-17: Token Exchange error messages leak details (I-09 T9) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Replace detailed error messages with generic 'invalid_grant' responses.\n\n"
                    "RR-18: Client Credentials no audience restriction (I-03 T14) \u2014 High \u2014 Open\n"
                    "  \u2192 Implement per-app audience whitelist for client_credentials grant.\n\n"
                    "RR-19: Missing security response headers globally (I-16 T5) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Add HSTS, X-Content-Type-Options, X-Frame-Options, CSP to all OAuth responses.\n\n"
                    "RR-20: No DB cleanup of expired auth codes/requests (I-10 T8) \u2014 High \u2014 Open\n"
                    "  \u2192 Implement periodic cleanup job for expired auth data. Zero cleanup exists currently.\n\n"
                    "RR-21: Refresh token JWT exposes sensitive claims (I-15 T9) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Consider opaque refresh tokens or minimal JWT claims.\n\n"
                    "RR-22: Auth code not checked for expiry in SQL (I-10 T9) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Add expiry check in SELECT query, not just in-memory.\n\n"
                    "RR-23: HTTP redirect URIs still allowed (I-01 T12) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Reject http:// redirect URIs entirely per RFC 9700 Section 2.6.\n\n"
                    "RR-24: Auth context consumed before assertion verification (I-02 T7) \u2014 High \u2014 Open\n"
                    "  \u2192 Verify assertion before consuming auth request context, or restore context on verification failure.\n\n"
                    "RR-25: Auth request flooding on authorize endpoint (I-01 T19) \u2014 Medium \u2014 Open\n"
                    "  \u2192 Rate limit /oauth2/authorize. Each request creates DB entries and flows.\n\n"
                    "RR-26: Missing at_hash in ID tokens (I-15 T8) \u2014 High \u2014 Open\n"
                    "  \u2192 Compute at_hash per OIDC Core Section 3.3.2.11 for access token substitution prevention.")
                break

    # ═════════════════════════════════════════════════════════════════════
    # 18. Document Lifecycle
    # ═════════════════════════════════════════════════════════════════════
    idx_dl = find_paragraph_index(doc, "The threat model moved to")
    if idx_dl >= 0:
        set_paragraph_text(doc.paragraphs[idx_dl], "To be moved to Security Review Documents.")
    idx_dl2 = find_paragraph_index(doc, "Threat model reviewed by security team")
    if idx_dl2 >= 0:
        set_paragraph_text(doc.paragraphs[idx_dl2], "To be reviewed by security team: TBD.")

    # ═════════════════════════════════════════════════════════════════════
    # 18b. Consultation Sessions — Replace template placeholders
    # ═════════════════════════════════════════════════════════════════════
    idx_cs = find_paragraph_index(doc, "Session 1:")
    if idx_cs >= 0:
        set_paragraph_text(doc.paragraphs[idx_cs], "No consultation sessions conducted yet.")
        # Remove the remaining template placeholder paragraphs (Session 1 fields + Session 2 block)
        for pi in range(idx_cs+1, min(idx_cs+18, len(doc.paragraphs))):
            p = doc.paragraphs[pi]
            t = p.text.strip()
            if t.startswith("Date:") or t.startswith("Participants") or \
               t.startswith("Session recording") or t.startswith("Notes") or \
               t.startswith("Action Items") or t.startswith("Session 2") or \
               t.startswith("[LINK]") or t == "":
                clear_paragraph(p)
            elif "Heading" in (p.style.name or ""):
                break  # Reached next section

    # ═════════════════════════════════════════════════════════════════════
    # 19. Appendix
    # ═════════════════════════════════════════════════════════════════════
    print("  Writing appendix...")
    idx_feat = find_paragraph_index(doc, "Feature/Product Documentation")
    if idx_feat >= 0:
        for pi in range(idx_feat+1, min(idx_feat+3, len(doc.paragraphs))):
            if doc.paragraphs[pi].text.strip() in ("<insert>",""):
                set_paragraph_text(doc.paragraphs[pi],
                    "\u2022 api/authentication.yaml\n"
                    "\u2022 api/flow-execution.yaml\n"
                    "\u2022 api/application.yaml\n"
                    "\u2022 api/design.yaml\n"
                    "\u2022 backend/internal/oauth/\n"
                    "\u2022 backend/internal/application/model/oauth_app.go")
                break

    idx_cnad = find_paragraph_index(doc, "CNAD/Application Development Checklist")
    if idx_cnad >= 0:
        for pi in range(idx_cnad+1, min(idx_cnad+3, len(doc.paragraphs))):
            if doc.paragraphs[pi].text.strip() in ("<insert>",""):
                set_paragraph_text(doc.paragraphs[pi],
                    "\u2022 OAuth endpoint validation: All inputs (redirect_uri, "
                    "scope, grant_type, response_type) are validated against "
                    "registered application configuration.\n"
                    "\u2022 Client authentication: client_secret_basic, "
                    "client_secret_post, and none methods supported with "
                    "constant-time comparison.\n"
                    "\u2022 Token lifecycle: Authorization codes are single-use, "
                    "configurable expiry. Tokens include iss, sub, aud, exp, iat, "
                    "jti claims.\n"
                    "\u2022 PKCE: Required for public clients. S256 and plain methods "
                    "supported. Downgrade prevention enforced.\n"
                    "\u2022 Error responses: Pre-redirect errors go to error page. "
                    "Post-redirect errors redirect with error+state. No sensitive "
                    "details in error responses.")
                break

    idx_conf = find_paragraph_index(doc, "Sample Configs:")
    if idx_conf >= 0:
        for pi in range(idx_conf+1, min(idx_conf+3, len(doc.paragraphs))):
            t = doc.paragraphs[pi].text.strip()
            if t.startswith("<insert") or t == "":
                set_paragraph_text(doc.paragraphs[pi],
                    "# deployment.yaml (OAuth-relevant)\n"
                    "server:\n  hostname: \"localhost\"\n  port: 8090\n\n"
                    "tls:\n  min_version: \"1.3\"\n  cert_file: \"...server.cert\"\n  key_file: \"...server.key\"\n\n"
                    "jwt:\n  preferred_key_id: \"default-key\"\n  validity_period: 3600\n  leeway: 0\n\n"
                    "oauth:\n  refresh_token:\n    renew_on_grant: false\n    validity_period: 86400\n"
                    "  authorization_code:\n    validity_period: 300\n"
                    "  dcr:\n    insecure: false  # MUST be false in production\n\n"
                    "# Per-app config (in application model):\n"
                    "# token.access_token.validity_period, token.access_token.user_attributes\n"
                    "# token.id_token.validity_period, token.id_token.user_attributes\n"
                    "# user_info.user_attributes, scope_claims, pkce_required, public_client")
                break

    idx_audit = find_paragraph_index(doc, "Sample Audit Logs:")
    if idx_audit >= 0:
        for pi in range(idx_audit+1, min(idx_audit+3, len(doc.paragraphs))):
            t = doc.paragraphs[pi].text.strip()
            if t.startswith("<insert") or t == "":
                set_paragraph_text(doc.paragraphs[pi],
                    "Recommended audit events:\n"
                    "\u2022 Token issuance: grant_type, client_id, scopes, sub, duration_ms\n"
                    "\u2022 Client auth failures: client_id, method, error\n"
                    "\u2022 Auth code gen/consumption: code_id, client_id, user_id\n"
                    "\u2022 DCR: client_id, redirect_uris, grants\n"
                    "\u2022 Introspection: caller, token_hint, active\n"
                    "\u2022 Observability: TOKEN_ISSUED, TOKEN_ISSUANCE_FAILED")
                break

    # ── RFC 9700 Compliance Table ──
    doc.add_heading("RFC 9700 Compliance Summary", level=3)
    rfc_t = doc.add_table(rows=1, cols=4)
    add_table_borders(rfc_t)
    set_table_cell_margins(rfc_t)
    for i, h in enumerate(["Requirement","Section","Status","Notes"]):
        set_cell_text(rfc_t.rows[0].cells[i], h, bold=True, size=TABLE_HEADER_SIZE)
        set_cell_shading(rfc_t.rows[0].cells[i], "D9E2F3")
    for r in [
        # Section 2.1 Protecting Redirect-Based Flows
        ("Exact redirect URI matching (MUST)",
         "2.1, 4.1.3",
         "Compliant",
         "Exact string comparison. No wildcards or pattern matching."),
        ("No open redirectors (MUST NOT)",
         "2.1, 4.11",
         "Compliant",
         "Invalid redirect \u2192 error page, not redirect to client."),
        ("CSRF protection (MUST)",
         "2.1, 4.7.1",
         "Compliant",
         "PKCE binds code to client instance. State param also supported."),
        ("Mix-up defense (REQUIRED when multi-AS)",
         "2.1, 4.4.2",
         "Compliant",
         "AS Metadata via /.well-known. iss claim in tokens per RFC 9207."),
        # Section 2.1.1 Authorization Code Grant
        ("Public clients MUST use PKCE",
         "2.1.1",
         "Compliant",
         "Enforced for public clients. RECOMMENDED for confidential."),
        ("AS MUST support PKCE",
         "2.1.1",
         "Partial",
         "S256 + plain methods. plain method should be disabled per Security BCP. Only S256 recommended."),
        ("AS MUST enforce code_verifier",
         "2.1.1",
         "Compliant",
         "Stored code_challenge verified at token endpoint."),
        ("PKCE downgrade prevention (MUST)",
         "2.1.1, 4.8.2",
         "Compliant",
         "code_verifier rejected if no code_challenge was in auth request."),
        # Section 2.1.2 Implicit Grant
        ("Implicit grant deprecated (SHOULD NOT)",
         "2.1.2",
         "Compliant",
         "Only response_type=code supported. No implicit/token."),
        # Section 2.2 Token Replay Prevention
        ("Sender-constrained access tokens (SHOULD)",
         "2.2.1",
         "Partial",
         "Bearer tokens only. DPoP (RFC 9449) / mTLS (RFC 8705) not yet implemented."),
        ("Refresh token replay detection (MUST for public)",
         "2.2.2, 4.14.2",
         "Compliant",
         "Refresh token rotation via renew_on_grant config."),
        # Section 2.3 Access Token Privilege Restriction
        ("Audience-restricted access tokens (SHOULD)",
         "2.3",
         "Compliant",
         "aud claim set. Resource indicators (RFC 8707) supported."),
        ("Access token scope restriction (SHOULD)",
         "2.3",
         "Non-Compliant",
         "Scope intersection on refresh only. Initial issuance: ValidateScopes is no-op. Any scope granted."),
        # Section 2.4 ROPC
        ("ROPC grant (MUST NOT)",
         "2.4",
         "Compliant",
         "Not implemented. No password grant handler."),
        # Section 2.5 Client Authentication
        ("Client authentication (SHOULD enforce)",
         "2.5",
         "Compliant",
         "client_secret_basic/post. Constant-time SHA-256 comparison."),
        ("Asymmetric client auth (RECOMMENDED)",
         "2.5",
         "Not Impl.",
         "Only symmetric (shared secret). private_key_jwt / mTLS not supported."),
        # Section 2.6 Other Recommendations
        ("AS Metadata (RECOMMENDED)",
         "2.6",
         "Compliant",
         "Both /.well-known/openid-configuration and /oauth-authorization-server."),
        ("No http redirect URIs (MUST NOT)",
         "2.6",
         "Non-Compliant",
         "http:// URIs allowed with warning flag. Should reject entirely per RFC 9700."),
        ("End-to-end TLS (RECOMMENDED)",
         "2.6",
         "Compliant",
         "TLS 1.3 minimum. Configurable cert/key in deployment.yaml."),
        ("CORS on authz endpoint (MUST NOT)",
         "2.6",
         "Non-Compliant",
         "CORS middleware applied to /oauth2/authorize. Should be removed."),
        ("client_id impersonation prevention (SHOULD NOT)",
         "2.6, 4.15.1",
         "N/A",
         "client_id is server-generated UUID. Not user-controllable."),
        # Section 4 Attacks and Mitigations
        ("Auth code single-use (MUST)",
         "4.2.4, 4.5",
         "Compliant",
         "Code set INACTIVE on first retrieval from DB."),
        ("No 307 redirect (MUST NOT)",
         "4.12",
         "Compliant",
         "Uses 302 Found for all authorization redirects."),
        ("Clickjacking prevention (MUST)",
         "4.16",
         "Not Impl.",
         "No X-Frame-Options or CSP frame-ancestors on authorization endpoint."),
        ("Referrer-Policy (SHOULD)",
         "4.2.4",
         "Not Impl.",
         "No Referrer-Policy header set on authorization/callback responses."),
        ("Refresh token scope binding (MUST)",
         "4.14.2",
         "Compliant",
         "Refresh tokens bound to original scope. Cannot escalate."),
        ("Access tokens not in URI query (MUST NOT)",
         "4.3.2",
         "Compliant",
         "Tokens only in Authorization header or POST body."),
        ("TLS proxy header sanitization (MUST)",
         "4.13",
         "N/A",
         "Deployment-specific. No reverse proxy in default config."),
        ("Token revocation (RECOMMENDED)",
         "2.2.2",
         "Not Impl.",
         "/oauth2/revoke constant defined but no handler registered. Stolen tokens irrevocable until expiry."),
        ("OIDC nonce parameter (RECOMMENDED for code flow)",
         "OIDC Core 3.1.2.1",
         "Not Impl.",
         "nonce not extracted, stored, or included in ID tokens. Enables ID token replay."),
        ("Request body size limits (SHOULD)",
         "General",
         "Not Impl.",
         "No http.MaxBytesReader on any endpoint. Unbounded POST bodies accepted."),
    ]:
        add_row(rfc_t, r)

    # ── Endpoint Registry Table ──
    doc.add_heading("Complete Endpoint Registry", level=3)
    ep_t = doc.add_table(rows=1, cols=4)
    add_table_borders(ep_t)
    set_table_cell_margins(ep_t)
    for i, h in enumerate(["Method","Route","Handler","Auth"]):
        set_cell_text(ep_t.rows[0].cells[i], h, bold=True, size=TABLE_HEADER_SIZE)
        set_cell_shading(ep_t.rows[0].cells[i], "D9E2F3")
    for r in [
        ("GET","/oauth2/authorize","HandleAuthorizeGetRequest","CORS"),
        ("POST","/oauth2/auth/callback","HandleAuthCallbackPostRequest","CORS"),
        ("POST","/oauth2/token","HandleTokenRequest","CORS + ClientAuth MW"),
        ("POST","/oauth2/introspect","HandleIntrospect","CORS only (GAP)"),
        ("GET|POST","/oauth2/userinfo","HandleUserInfo","CORS + Bearer"),
        ("GET","/oauth2/jwks","HandleJWKSRequest","CORS"),
        ("GET","/.well-known/oauth-authorization-server","HandleOAuth2ASMetadata","CORS"),
        ("GET","/.well-known/openid-configuration","HandleOIDCDiscovery","CORS"),
        ("POST","/oauth2/dcr/register","HandleDCRRegistration","CORS + perm check"),
        ("\u2014","/oauth2/revoke (NOT IMPL.)","\u2014","Constant only"),
        ("\u2014","/oauth2/logout (NOT IMPL.)","\u2014","Constant only"),
    ]:
        add_row(ep_t, r)

    # ═════════════════════════════════════════════════════════════════════
    # POST-PROCESSING: Apply spacing to any remaining body paragraphs
    # and cell margins to any untouched tables
    # ═════════════════════════════════════════════════════════════════════
    print("  Applying final spacing pass...")
    heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name in heading_styles:
            continue  # Keep heading spacing from styles
        if not p.text.strip():
            continue  # Skip empty paragraphs
        if p.paragraph_format.line_spacing is None:
            p.paragraph_format.line_spacing = LINE_SPACING
        if p.paragraph_format.space_after is None:
            p.paragraph_format.space_after = SPACE_AFTER_BODY

    # Apply cell margins to any table that doesn't have them yet
    for t in doc.tables:
        tblPr = t._tbl.find(qn('w:tblPr'))
        if tblPr is not None and tblPr.find(qn('w:tblCellMar')) is None:
            set_table_cell_margins(t)

    # ═════════════════════════════════════════════════════════════════════
    # SAVE
    # ═════════════════════════════════════════════════════════════════════
    doc.save(OUTPUT_PATH)
    print(f"\n\u2713 Generated: {OUTPUT_PATH}")
    print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
