#!/usr/bin/env python3
"""Compare reference App Gateway threat model with our Thunder OAuth2 threat model."""
from docx import Document

ref = Document('Threat Model - App gateway support for web app authentication.docx')
our = Document('[Threat Model] - Thunder OAuth2.docx')

print('=' * 80)
print('REFERENCE DOC - Full heading hierarchy')
print('=' * 80)
for i, p in enumerate(ref.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f'  [{i}] {p.style.name}: {p.text[:120]}')

print()
print('=' * 80)
print('REFERENCE DOC - Table details')
print('=' * 80)
for ti, t in enumerate(ref.tables):
    headers = [c.text.strip()[:40] for c in t.rows[0].cells] if t.rows else []
    sample = ''
    if len(t.rows) > 1:
        sample = ' | '.join([c.text.strip()[:40] for c in t.rows[1].cells])
    print(f'  T{ti}: rows={len(t.rows)}, headers={headers}')
    if sample:
        print(f'       sample: {sample[:180]}')

print()
print('=' * 80)
print('OUR DOC - Heading hierarchy')
print('=' * 80)
for i, p in enumerate(our.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f'  [{i}] {p.style.name}: {p.text[:120]}')

print()
print('=' * 80)
print('OUR DOC - Table details')
print('=' * 80)
for ti, t in enumerate(our.tables):
    headers = [c.text.strip()[:40] for c in t.rows[0].cells] if t.rows else []
    sample = ''
    if len(t.rows) > 1:
        sample = ' | '.join([c.text.strip()[:40] for c in t.rows[1].cells])
    print(f'  T{ti}: rows={len(t.rows)}, headers={headers}')
    if sample:
        print(f'       sample: {sample[:180]}')

# Now extract per-interaction content from reference doc to see what subsections each interaction has
print()
print('=' * 80)
print('REFERENCE DOC - Per-interaction content structure')
print('=' * 80)
in_interaction = False
current_interaction = None
for i, p in enumerate(ref.paragraphs):
    if p.style.name == 'Heading 3' and 'Interaction' in p.text:
        current_interaction = p.text[:80]
        in_interaction = True
        print(f'\n--- {current_interaction} ---')
    elif p.style.name.startswith('Heading') and in_interaction:
        if p.style.name in ('Heading 1', 'Heading 2'):
            in_interaction = False
        else:
            print(f'  {p.style.name}: {p.text[:100]}')
    elif in_interaction and p.text.strip():
        txt = p.text.strip()[:100]
        style = p.style.name
        if 'Bold' in style or any(kw in txt.lower() for kw in ['assets involved', 'description', 'access control', 'security consideration', 'cookie']):
            print(f'  ** {style}: {txt}')
        elif len(txt) > 5:
            print(f'     {style}: {txt}')

# Now extract per-interaction content from our doc 
print()
print('=' * 80)
print('OUR DOC - Per-interaction content structure')
print('=' * 80)
in_interaction = False
current_interaction = None
for i, p in enumerate(our.paragraphs):
    if p.style.name == 'Heading 3' and '[I-' in p.text:
        current_interaction = p.text[:80]
        in_interaction = True
        print(f'\n--- {current_interaction} ---')
    elif p.style.name.startswith('Heading') and in_interaction:
        if p.style.name in ('Heading 1', 'Heading 2'):
            in_interaction = False
        else:
            print(f'  {p.style.name}: {p.text[:100]}')
    elif in_interaction and p.text.strip():
        txt = p.text.strip()[:120]
        style = p.style.name
        print(f'     {style}: {txt}')
