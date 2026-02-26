#!/usr/bin/env python3
"""Extract detailed per-interaction structure from reference doc."""
from docx import Document

ref = Document('Threat Model - App gateway support for web app authentication.docx')

# Find the "Threats and Mitigations" section and extract all content within interactions
print("=" * 80)
print("REFERENCE DOC - Detailed per-interaction structure (Interaction 02: Login)")
print("=" * 80)

# Find para index range for Interaction 02
start_idx = None
end_idx = None
for i, p in enumerate(ref.paragraphs):
    if 'Interactions [02]: Login' in p.text and p.style.name == 'Heading 3':
        start_idx = i
    elif start_idx and p.style.name in ('Heading 3', 'Heading 2', 'Heading 1') and i > start_idx + 1:
        end_idx = i
        break

if start_idx and end_idx:
    print(f"Paragraphs {start_idx} to {end_idx}")
    for i in range(start_idx, min(end_idx, start_idx + 80)):
        p = ref.paragraphs[i]
        style = p.style.name
        text = p.text.strip()
        # Check for bold runs
        bold_runs = [r.text for r in p.runs if r.bold]
        bold_marker = " [BOLD]" if bold_runs else ""
        if text:
            print(f"  [{i}] {style}{bold_marker}: {text[:150]}")
        elif style.startswith('Heading'):
            print(f"  [{i}] {style}: (empty)")

# Now find tables that appear between these paragraphs
# docx doesn't directly map tables to paragraph positions, so let's track by counting body elements
print()
print("=" * 80)
print("REFERENCE DOC - Body element ordering for Interaction 02")
print("=" * 80)

from docx.oxml.ns import qn
body = ref.element.body
para_idx = 0
table_idx = 0
in_range = False
for child in body:
    if child.tag == qn('w:p'):
        if para_idx == start_idx:
            in_range = True
        if para_idx == end_idx:
            in_range = False
        if in_range:
            text = child.text if hasattr(child, 'text') else ''
            # Get full text from all runs
            full_text = ''.join(node.text or '' for node in child.iter(qn('w:t')))
            print(f"  PARA [{para_idx}]: {full_text[:120]}")
        para_idx += 1
    elif child.tag == qn('w:tbl'):
        if in_range:
            # Get table headers
            first_row = child.find(qn('w:tr'))
            if first_row is not None:
                cells = first_row.findall(qn('w:tc'))
                headers = []
                for c in cells:
                    cell_text = ''.join(node.text or '' for node in c.iter(qn('w:t')))
                    headers.append(cell_text.strip()[:30])
                # Count rows
                rows = child.findall(qn('w:tr'))
                print(f"  TABLE [T{table_idx}]: rows={len(rows)}, headers={headers}")
        table_idx += 1

print()
print("=" * 80)  
print("REFERENCE DOC - Body element ordering for Interaction 01 (Config resolution)")
print("=" * 80)

# Find para index range for Interaction 01 in Threats section
start_idx_01 = None
end_idx_01 = None
for i, p in enumerate(ref.paragraphs):
    if 'Interactions [01]: Config resolution' in p.text and p.style.name == 'Heading 3' and i > 120:
        start_idx_01 = i
    elif start_idx_01 and p.style.name == 'Heading 3' and i > start_idx_01 + 1:
        end_idx_01 = i
        break

if start_idx_01 and end_idx_01:
    para_idx = 0
    table_idx = 0
    in_range = False
    for child in body:
        if child.tag == qn('w:p'):
            if para_idx == start_idx_01:
                in_range = True
            if para_idx == end_idx_01:
                in_range = False
            if in_range:
                full_text = ''.join(node.text or '' for node in child.iter(qn('w:t')))
                style_elem = child.find(qn('w:pPr'))
                pstyle = ''
                if style_elem is not None:
                    style_name = style_elem.find(qn('w:pStyle'))
                    if style_name is not None:
                        pstyle = style_name.get(qn('w:val'), '')
                print(f"  PARA [{para_idx}] ({pstyle}): {full_text[:150]}")
            para_idx += 1
        elif child.tag == qn('w:tbl'):
            if in_range:
                first_row = child.find(qn('w:tr'))
                if first_row:
                    cells = first_row.findall(qn('w:tc'))
                    headers = []
                    for c in cells:
                        cell_text = ''.join(node.text or '' for node in c.iter(qn('w:t')))
                        headers.append(cell_text.strip()[:30])
                    rows = child.findall(qn('w:tr'))
                    print(f"  TABLE [T{table_idx}]: rows={len(rows)}, headers={headers}")
            table_idx += 1


# Now check our doc's interaction structure for comparison
print()
print("=" * 80)
print("OUR DOC - Body element ordering for I-01 (Authorization Request)")
print("=" * 80)

our = Document('[Threat Model] - Thunder OAuth2.docx')
# Find I-01 heading
start_our = None
end_our = None
for i, p in enumerate(our.paragraphs):
    if '[I-01]' in p.text and p.style.name == 'Heading 3':
        start_our = i
    elif start_our and p.style.name in ('Heading 3', 'Heading 2', 'Heading 1') and i > start_our + 1:
        end_our = i
        break

if start_our and end_our:
    body_our = our.element.body
    para_idx = 0
    table_idx = 0
    in_range = False
    for child in body_our:
        if child.tag == qn('w:p'):
            if para_idx == start_our:
                in_range = True
            if para_idx == end_our:
                in_range = False
            if in_range:
                full_text = ''.join(node.text or '' for node in child.iter(qn('w:t')))
                style_elem = child.find(qn('w:pPr'))
                pstyle = ''
                if style_elem is not None:
                    style_name = style_elem.find(qn('w:pStyle'))
                    if style_name is not None:
                        pstyle = style_name.get(qn('w:val'), '')
                if full_text.strip():
                    print(f"  PARA [{para_idx}] ({pstyle}): {full_text[:150]}")
            para_idx += 1
        elif child.tag == qn('w:tbl'):
            if in_range:
                first_row = child.find(qn('w:tr'))
                if first_row:
                    cells = first_row.findall(qn('w:tc'))
                    headers = []
                    for c in cells:
                        cell_text = ''.join(node.text or '' for node in c.iter(qn('w:t')))
                        headers.append(cell_text.strip()[:30])
                    rows = child.findall(qn('w:tr'))
                    print(f"  TABLE [T{table_idx}]: rows={len(rows)}, headers={headers}")
            table_idx += 1
